import os
import inspect
import json
import logging
import time
import faiss
from google import genai
from google.genai import types
import pdfplumber
import numpy as np
import re
from typing import List, Dict, Any, Tuple, Optional, AsyncIterator
from sentence_transformers import SentenceTransformer
from ratelimit import limits, sleep_and_retry
import concurrent.futures
from datetime import datetime
import hashlib
from functools import lru_cache
import asyncio
from sklearn.metrics.pairwise import cosine_similarity
import docx
import PyPDF2
from pathlib import Path
from docx import Document
import random
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuration
_GEMINI_KEYS = []
_ACTIVE_GEMINI_INDEX: Optional[int] = None
# Track exhausted keys with cooldown timestamps: {key_index: cooldown_until_timestamp}
_EXHAUSTED_KEYS: Dict[int, float] = {}
_EXHAUSTION_COOLDOWN_SECONDS = 300  # 5 minutes cooldown for exhausted keys
# Track permanently invalid keys (leaked, revoked, etc.): {key_index: reason}
_INVALID_KEYS: Dict[int, str] = {}
# Track permanently invalid keys (leaked, revoked, etc.): {key_index: reason}
_INVALID_KEYS: Dict[int, str] = {}

def _reload_gemini_keys():
    """Reload Gemini API keys from environment variables."""
    global _GEMINI_KEYS, _ACTIVE_GEMINI_INDEX
    _GEMINI_KEYS = []
    _ACTIVE_GEMINI_INDEX = None
    
    # Reload .env file if it exists
    env_path = Path(__file__).parent.parent / ".env"
    if env_path.exists():
        from dotenv import load_dotenv
        load_dotenv(env_path, override=True)
    else:
        from dotenv import load_dotenv
        load_dotenv(override=True)
    
    for key_candidate in (
        os.getenv("GOOGLE_API_KEY1"),
        os.getenv("GOOGLE_API_KEY2"),
        os.getenv("GOOGLE_API_KEY3"),
        os.getenv("GOOGLE_API_KEY4"),
    ):
        if key_candidate:
            value = key_candidate.strip()
            if value and value not in _GEMINI_KEYS:
                _GEMINI_KEYS.append(value)
    
    logger.info(f"✅ Loaded {len(_GEMINI_KEYS)} Gemini API key(s) from environment")
    if _GEMINI_KEYS:
        # Show first 10 chars of each key for verification
        key_previews = [f"KEY{i+1}: {key[:10]}..." for i, key in enumerate(_GEMINI_KEYS)]
        logger.info(f"   Keys: {', '.join(key_previews)}")
        logger.info(f"   Note: If keys are from different Google Cloud projects/accounts, each will have separate quotas")
    else:
        logger.warning("⚠️  No Gemini API keys found in environment variables!")

# Initial load
_reload_gemini_keys()

# Initialize the embedding model
logger.info("Initializing embedding model...")
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Directory structure for embeddings and indexes
EMBEDDING_DIR = "embeddings"
INDEX_DIR = "faiss_indexes"
CACHE_DIR = "compliance_cache"
PDF_FOLDER = "compliance_frameworks"

EMBEDDING_CACHE_FILE = os.path.join(EMBEDDING_DIR, "document_embeddings.npy")
DOCUMENT_MAP_FILE = os.path.join(EMBEDDING_DIR, "document_map.json")
FAISS_INDEX_FILE = os.path.join(INDEX_DIR, "compliance_index.faiss")
QUERY_CACHE_FILE = os.path.join(CACHE_DIR, "query_cache.json")

# Create directories if they don't exist
for directory in [EMBEDDING_DIR, INDEX_DIR, CACHE_DIR]:
    if not os.path.exists(directory):
        os.makedirs(directory)

# Initialize query cache
QUERY_CACHE = {}
if os.path.exists(QUERY_CACHE_FILE):
    try:
        with open(QUERY_CACHE_FILE, 'r', encoding='utf-8') as f:
            QUERY_CACHE = json.load(f)
        logger.info(f"✅ Loaded {len(QUERY_CACHE)} cached queries from {QUERY_CACHE_FILE}")
    except Exception as e:
        logger.error(f"❌ Error loading query cache: {e}")
        QUERY_CACHE = {}  # Start with empty cache if loading fails
else:
    logger.info(f"📝 No existing cache file found. Starting with empty cache.")

def save_query_cache():
    """Save query cache to disk immediately"""
    try:
        # Create cache directory if it doesn't exist
        os.makedirs(CACHE_DIR, exist_ok=True)
        
        # Save with proper formatting
        with open(QUERY_CACHE_FILE, 'w', encoding='utf-8') as f:
            json.dump(QUERY_CACHE, f, indent=2, ensure_ascii=False)
        logger.info(f"💾 Saved {len(QUERY_CACHE)} queries to cache file: {QUERY_CACHE_FILE}")
    except Exception as e:
        logger.error(f"❌ Error saving query cache: {e}")
        raise

def classify_document_type(text: str, allow_general_docs: bool = False) -> str:
    """Classify uploaded document content with hybrid rules + LLM fallback.

    Returns:
        - 'privacy_policy' for privacy policies
        - 'terms_and_conditions' for terms of service
        - 'general_documentation' for system/technical docs (when allow_general_docs=True)
        - 'other' for unrelated documents (CV, personal docs, academic content, etc.)
    """
    if not text:
        return "other"
    t = (text or "").lower()
    try:
        logger.info(f"DocType: start len={len(t)} preview='{t[:120].strip()}'")
    except Exception:
        pass

    # ISO/Compliance Framework documents (CHECK FIRST - highest priority rejection)
    # These are the actual compliance standards/frameworks themselves, not user policies
    compliance_framework_indicators = [
        "iso/iec 27001", "iso 27001", "iso27001", "iso 27002", "iso27002", "iso/iec", 
        "information technology  — security",  # Exact match from upload preview
        "information security management systems — requirements",
        "information security management system requirements",
        "security techniques — informati",  # Catches partial from upload
        "security techniques — information security",
        "security techniques information security",
        "technologies de l'information", "tecnolog", "technologies d'",  # ISO docs often have multiple languages
        "international standard", "this international standard",
        "foreword", "normative references",  # ISO structure markers
        "soc 2 controls", "soc 2 trust services criteria", "trust services criteria",
        "nist framework", "nist cybersecurity framework", "nist sp 800", "special publication",
        "pci dss standard", "pci security standards", "payment card industry data security standard",
        "hipaa regulation", "hipaa administrative simplification", "health insurance portability",
        "annex a controls", "annex a.", "annex b.", "control objective", "control category",
        "compliance standard", "regulatory framework", "certification requirements",
        "audit criteria", "control framework", "maturity model",
        "iso iec", "iec iso", "first edition", "second edition", "third edition"  # ISO edition markers
    ]
    
    # Check if this is a compliance framework document itself - REJECT IMMEDIATELY
    framework_hits = sum(1 for k in compliance_framework_indicators if k in t)
    if framework_hits >= 1:
        logger.info(f"DocType: Detected compliance framework document (hits: {framework_hits}) - REJECTING")
        return "other"
    
    # Direct phrase shortcuts for high precision (but AFTER framework check)
    try:
        contains_privacy = "privacy policy" in t
        # More specific terms matching to avoid false positives with "terms and definitions" etc.
        contains_terms = (
            ("terms and conditions" in t) or 
            ("terms of service" in t) or 
            ("terms of use" in t) or
            ("user agreement" in t and "liability" in t)  # Terms docs always have liability clauses
        )
        logger.info(
            f"DocType: phrase checks -> privacy='{'yes' if contains_privacy else 'no'}', "
            f"terms='{'yes' if contains_terms else 'no'}'"
        )
        if contains_privacy:
            return "privacy_policy"
        if contains_terms:
            return "terms_and_conditions"
    except Exception:
        pass
    
    # Quick exits for CV/resume keywords
    cv_hits = 0
    for kw in [
        "curriculum vitae", "resume", "linkedin.com/in", "objective", "experience",
        "education", "skills", "certifications", "projects", "work history", "references"
    ]:
        if kw in t:
            cv_hits += 1
    if cv_hits >= 3:
        return "other"
    
    # Academic/Educational content indicators (quizzes, exams, assignments, lectures)
    academic_indicators = [
        "quiz", "exam", "test", "midterm", "final exam", "assignment", "homework",
        "solution:", "question", "marks:", "total marks", "reg. no:", "registration number",
        "semester", "course code", "clo", "plo", "learning outcome", "lecture notes",
        "department of", "university", "college", "instructor:", "professor:",
        "due date:", "submission", "grading rubric", "answer key", "correct answer"
    ]
    
    academic_hits = sum(1 for k in academic_indicators if k in t)
    if academic_hits >= 3:
        return "other"
    
    # Privacy policy indicators (expanded and more flexible)
    privacy_indicators = [
        "privacy policy", "data privacy", "personal data", "data controller", "data processor",
        "gdpr", "ccpa", "cpra", "data subject", "right to access", "right to erasure",
        "how we collect", "how we use", "retention", "cookie policy", "consent",
        "information we collect", "what information", "data we collect", "personal information",
        "data practices", "privacy notice", "data protection", "user data", "customer data",
        "third parties", "third-party", "data sharing", "information sharing",
        "opt-out", "opt out", "marketing communications", "analytics", "tracking"
    ]
    
    # Terms & conditions indicators
    terms_indicators = [
        "terms and conditions", "terms of service", "terms of use", "user agreement",
        "limitation of liability", "governing law", "arbitration", "indemnification",
        "warranty disclaimer", "acceptable use", "license grant", "intellectual property",
        "user obligations", "termination", "suspension of account"
    ]
    
    # General documentation indicators (REAL system/technical docs ONLY)
    # These must indicate actual system documentation, not academic content
    general_doc_indicators = [
        "system architecture document", "technical design document", "api specification",
        "software architecture", "deployment guide", "infrastructure setup",
        "authentication implementation", "authorization mechanism", "security architecture",
        "data flow diagram", "database design", "microservices architecture",
        "production deployment", "development environment", "staging environment",
        "disaster recovery plan", "backup strategy", "monitoring setup",
        "ci/cd pipeline", "devops workflow", "release notes"
    ]
    
    # Count matches for each category
    privacy_matches = [k for k in privacy_indicators if k in t]
    terms_matches = [k for k in terms_indicators if k in t]
    general_matches = [k for k in general_doc_indicators if k in t] if allow_general_docs else []
    
    p_score = len(privacy_matches)
    t_score = len(terms_matches)
    g_score = len(general_matches)
    
    # Log classification details for debugging
    logger.info(
        f"DocType scores - Privacy: {p_score} (matches: {privacy_matches[:3]}), "
        f"Terms: {t_score}, General: {g_score}, Academic: {academic_hits}"
    )
    
    # Prefer privacy/terms with >=1 strong indicator (LLM will handle borderline cases)
    if p_score >= 1 and p_score >= max(t_score, g_score):
        return "privacy_policy"
    if t_score >= 1 and t_score >= max(p_score, g_score):
        return "terms_and_conditions"
    
    # If general docs allowed and detected (and NOT academic content)
    if allow_general_docs and g_score >= 2 and academic_hits < 3:
        return "general_documentation"
    
    # LLM fallback for ambiguous cases
    try:
        logger.info(
            f"DocType: LLM fallback (privacy={p_score}, terms={t_score}, general={g_score}, academic={academic_hits}, allow_general_docs={allow_general_docs})"
        )
        prompt = (
            "Classify the following document into ONE category: \n"
            "- privacy_policy \n"
            "- terms_and_conditions \n"
            + ("- general_documentation \n" if allow_general_docs else "") +
            "- other \n\n"
            "Instructions: base your decision on purpose and structure, not isolated keywords. "
            "Templates with placeholders like [Your Company Name] are valid. "
            "Academic content (quizzes, exams), personal documents (CVs), compliance framework documents (ISO standards, SOC 2 controls, NIST frameworks), and regulatory documents are 'other'.\n"
            "IMPORTANT: If the document IS the compliance standard/framework itself (like ISO 27001 standard document), classify as 'other'.\n"
            "Only user-created policies and documentation should be privacy_policy, terms_and_conditions, or general_documentation.\n\n"
            f"Document (first 2000 chars):\n{text[:2000]}\n\n"
            "Respond with ONLY one of: privacy_policy, terms_and_conditions, "
            + ("general_documentation, " if allow_general_docs else "") + "other"
        )
        llm_response = rate_limited_generate_content(prompt, temperature=0.1, max_tokens=20)
        logger.info(f"DocType: LLM raw='{str(llm_response)[:200]}'")
        classification = (llm_response or "").strip().lower()
        valid = ["privacy_policy", "terms_and_conditions", "general_documentation", "other"]
        if not allow_general_docs and "general_documentation" in valid:
            valid.remove("general_documentation")
        if classification in valid:
            logger.info(f"DocType: LLM classified='{classification}'")
            return classification
        logger.warning(f"DocType: invalid LLM classification '{classification}', defaulting to 'other'")
        return "other"
    except Exception as e:
        logger.warning(f"DocType: LLM classification failed: {e}", exc_info=True)
    return "other"

# Set up the Gemini model
MODEL_ID = "gemini-3-flash-preview"
client = None


def _configure_model_for_index(index: int) -> bool:
    global client, _ACTIVE_GEMINI_INDEX
    if index < 0 or index >= len(_GEMINI_KEYS):
        return False
    
    # Never configure invalid keys
    if _is_key_invalid(index):
        logger.debug(f"Refusing to configure invalid key #{index + 1}")
        return False
    
    key = _GEMINI_KEYS[index]
    if not key:
        return False
    try:
        client = genai.Client(api_key=key)
        _ACTIVE_GEMINI_INDEX = index
        key_preview = key[:10] + "..." if len(key) > 10 else key
        logger.info("Gemini configured successfully with key #%d (%s)", index + 1, key_preview)
        return True
    except Exception as exc:
        logger.warning("Failed to configure Gemini with key #%d: %s", index + 1, exc)
        client = None
        return False


def _ensure_model_initialized() -> bool:
    global client, _ACTIVE_GEMINI_INDEX
    if client is not None:
        # Check if current key is still valid
        if _ACTIVE_GEMINI_INDEX is not None and _is_key_invalid(_ACTIVE_GEMINI_INDEX):
            logger.warning(f"Current key #{_ACTIVE_GEMINI_INDEX + 1} is invalid, reinitializing...")
            client = None
            _ACTIVE_GEMINI_INDEX = None
    
    if client is not None:
        return True
    
    # Try to initialize with a valid, non-exhausted key
    for idx in range(len(_GEMINI_KEYS)):
        # Skip invalid keys
        if _is_key_invalid(idx):
            logger.debug(f"Skipping invalid key #{idx + 1} during initialization")
            continue
        # Skip exhausted keys (in cooldown)
        if _is_key_exhausted(idx):
            logger.debug(f"Skipping exhausted key #{idx + 1} during initialization")
            continue
        if _configure_model_for_index(idx):
            return True
    
    # If all non-exhausted keys failed, try exhausted keys (cooldown might have expired)
    logger.warning("All non-exhausted keys failed, trying exhausted keys...")
    for idx in range(len(_GEMINI_KEYS)):
        if _is_key_invalid(idx):
            continue
        if _configure_model_for_index(idx):
            logger.info(f"Initialized with exhausted key #{idx + 1} (cooldown expired)")
            return True
    
    if not _GEMINI_KEYS:
        logger.warning("Gemini not configured: no API key available (expected GOOGLE_API_KEY1 / GOOGLE_API_KEY2 / GOOGLE_API_KEY3 / GOOGLE_API_KEY4)")
    else:
        invalid_count = len(_INVALID_KEYS)
        exhausted_count = len([k for k in _EXHAUSTED_KEYS if time.time() < _EXHAUSTED_KEYS[k]])
        logger.warning(f"Gemini not configured: {invalid_count} invalid key(s), {exhausted_count} exhausted key(s), all keys unavailable")
    return False


def _is_key_invalid(key_index: int) -> bool:
    """Check if a key is permanently invalid (leaked, revoked, etc.)."""
    return key_index in _INVALID_KEYS

def _is_key_exhausted(key_index: int) -> bool:
    """Check if a key is currently in cooldown due to exhaustion."""
    # Don't check exhausted status for invalid keys - they're permanently disabled
    if _is_key_invalid(key_index):
        return True  # Treat invalid keys as exhausted so they're skipped
    
    if key_index not in _EXHAUSTED_KEYS:
        return False
    cooldown_until = _EXHAUSTED_KEYS[key_index]
    if time.time() < cooldown_until:
        return True
    # Cooldown expired, remove from exhausted list
    del _EXHAUSTED_KEYS[key_index]
    logger.info(f"Key #{key_index + 1} cooldown expired, available for use again")
    return False

def _mark_key_exhausted(key_index: int, cooldown_seconds: int = None):
    """Mark a key as exhausted and set cooldown period."""
    if cooldown_seconds is None:
        cooldown_seconds = _EXHAUSTION_COOLDOWN_SECONDS
    _EXHAUSTED_KEYS[key_index] = time.time() + cooldown_seconds
    logger.warning(f"Key #{key_index + 1} marked as exhausted, cooldown for {cooldown_seconds}s")

def _switch_to_fallback_key(skip_exhausted: bool = True) -> bool:
    """Switch to next available key, optionally skipping exhausted keys."""
    # Use the improved rotation function which handles invalid keys properly
    return _try_all_keys_rotation()

def _try_all_keys_rotation() -> bool:
    """Try rotating through all available keys systematically."""
    if not _GEMINI_KEYS:
        return False
    
    current = _ACTIVE_GEMINI_INDEX if _ACTIVE_GEMINI_INDEX is not None else -1
    
    # Try all keys in order, starting from next key
    for offset in range(1, len(_GEMINI_KEYS) + 1):
        idx = (current + offset) % len(_GEMINI_KEYS)
        # Skip if invalid (permanently disabled)
        if _is_key_invalid(idx):
            logger.debug(f"Skipping invalid key #{idx + 1} in rotation ({_INVALID_KEYS.get(idx, 'unknown reason')})")
            continue
        # Skip if exhausted (in cooldown)
        if _is_key_exhausted(idx):
            logger.debug(f"Skipping exhausted key #{idx + 1} in rotation (cooldown active)")
            continue
        if _configure_model_for_index(idx):
            logger.info(f"Rotated to key #{idx + 1} (attempt {offset}/{len(_GEMINI_KEYS)})")
            return True
    
    # If all are exhausted (but not invalid), try anyway after checking cooldown
    logger.warning("All non-exhausted keys failed, checking if any exhausted keys are ready...")
    for offset in range(1, len(_GEMINI_KEYS) + 1):
        idx = (current + offset) % len(_GEMINI_KEYS)
        # Never try invalid keys
        if _is_key_invalid(idx):
            continue
        # Only try exhausted keys if cooldown has expired
        if idx in _EXHAUSTED_KEYS:
            cooldown_until = _EXHAUSTED_KEYS[idx]
            if time.time() < cooldown_until:
                remaining = int(cooldown_until - time.time())
                logger.debug(f"Key #{idx + 1} still in cooldown ({remaining}s remaining)")
                continue
        if _configure_model_for_index(idx):
            logger.info(f"Rotated to key #{idx + 1} (cooldown expired)")
            return True
    
    return False


# Initialize model with first available key
if _ensure_model_initialized():
    logger.info(f"✅ Initialized Gemini model with key #{_ACTIVE_GEMINI_INDEX + 1 if _ACTIVE_GEMINI_INDEX is not None else 'unknown'}")
else:
    logger.warning("⚠️  Failed to initialize Gemini model with any available key")

# Optimized rate limiting configuration (env-driven)
CALLS_PER_MINUTE = int(os.getenv("GEMINI_CALLS_PER_MINUTE", "40"))
DELAY_BETWEEN_CALLS = float(os.getenv("LLM_CALL_SPREAD_SECONDS", "1.5"))
LLM_CONCURRENCY = int(os.getenv("LLM_CONCURRENCY", "2"))
last_call_time = 0
GEMINI_CALL_COUNT = 0

def _log_gemini_api_call(context: str, prompt: str, temperature: float, max_tokens: int, optimized: bool) -> None:
    """Centralized logging for every Gemini API call."""
    try:
        global GEMINI_CALL_COUNT
        GEMINI_CALL_COUNT += 1
        prompt_len = len(prompt) if isinstance(prompt, str) else 0
        prompt_digest = hash_text(prompt[:512]) if isinstance(prompt, str) else ""
        logger.info(
            f"GEMINI_CALL #{GEMINI_CALL_COUNT} optimized={optimized} context='{context}' prompt_len={prompt_len} temp={temperature} max_tokens={max_tokens} prompt_hash={prompt_digest}"
        )
    except Exception:
        pass

# Concurrency limiter and in-flight deduplication for Gemini calls
try:
    import threading
    _llm_semaphore = threading.Semaphore(LLM_CONCURRENCY)
    _inflight_lock = threading.Lock()
    _inflight_prompts = {}
except Exception:
    _llm_semaphore = None
    _inflight_lock = None
    _inflight_prompts = {}

# Define timing decorator first
def timing_decorator(func):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        execution_time = end_time - start_time
        logger.info(f"\n{func.__name__} execution time: {execution_time:.2f} seconds")
        return result
    return wrapper

def wait_for_rate_limit_optimized():
    """Optimized rate limiting with shorter delays."""
    global last_call_time
    current_time = time.time()
    time_since_last_call = current_time - last_call_time
    if time_since_last_call < DELAY_BETWEEN_CALLS:
        sleep_time = DELAY_BETWEEN_CALLS - time_since_last_call
        time.sleep(sleep_time)
    last_call_time = time.time()

def _wait_for_inflight(cache_key: str):
    """Block if the same prompt is already being processed."""
    if _inflight_lock is None:
        return None
    while True:
        with _inflight_lock:
            evt = _inflight_prompts.get(cache_key)
            if evt is None:
                evt = threading.Event()
                _inflight_prompts[cache_key] = evt
                return evt
        evt.wait(timeout=5)
        # After wait, re-check the map

def _extract_retry_delay(error_message: str) -> Optional[int]:
    """Extract retry delay from Gemini API error message."""
    if not error_message:
        return None
    
    # Look for "Please retry in X.XXs" pattern
    match = re.search(r'retry in ([\d.]+)s', error_message, re.IGNORECASE)
    if match:
        return int(float(match.group(1)) + 1)  # Add 1 second buffer
    
    # Look for retry_delay seconds in structured error
    match = re.search(r'retry_delay\s*{\s*seconds:\s*(\d+)', error_message, re.IGNORECASE)
    if match:
        return int(match.group(1)) + 1
    
    # Look for "Please retry in X.XX seconds"
    match = re.search(r'retry in ([\d.]+)\s*seconds?', error_message, re.IGNORECASE)
    if match:
        return int(float(match.group(1)) + 1)
    
    return None

def _is_quota_exhausted(error_message: str) -> bool:
    """Check if error indicates quota exhaustion (not just rate limit)."""
    if not error_message:
        return False
    
    error_lower = error_message.lower()
    quota_indicators = [
        "quota exceeded",
        "free_tier_requests",
        "limit: 0",
        "check your plan and billing",
        "exceeded your current quota",
        "quota_limit_value"
    ]
    
    # Check for RATE_LIMIT_EXCEEDED with quota_limit_value: "0" (this is quota exhaustion, not rate limit)
    if "rate_limit_exceeded" in error_lower and "quota_limit_value" in error_lower:
        if '"0"' in error_message or "value: \"0\"" in error_message or "quota_limit_value\" value: \"0\"" in error_message:
            return True
    
    return any(indicator in error_lower for indicator in quota_indicators)

def _is_key_leaked_or_invalid(error_message: str) -> bool:
    """Check if error indicates key is leaked, revoked, or permanently invalid."""
    if not error_message:
        return False
    
    error_lower = error_message.lower()
    invalid_indicators = [
        "api key was reported as leaked",
        "api key was leaked",
        "leaked",
        "revoked",
        "invalid api key",
        "api key not valid"
    ]
    
    # Only return True if it's clearly a key issue, not a general 403
    if "403" in error_message:
        if any(indicator in error_lower for indicator in ["leaked", "revoked", "invalid", "api key"]):
            return True
    
    return any(indicator in error_lower for indicator in invalid_indicators)

def _mark_key_invalid(key_index: int, reason: str):
    """Mark a key as permanently invalid (leaked, revoked, etc.)."""
    _INVALID_KEYS[key_index] = reason
    logger.error(f"Key #{key_index + 1} marked as PERMANENTLY INVALID: {reason}")

@sleep_and_retry
@limits(calls=CALLS_PER_MINUTE, period=60)
@timing_decorator
def rate_limited_generate_content_optimized(prompt: str, temperature: float = 0.1, max_tokens: int = 3200) -> str:
    """Optimized rate-limited content generation with reduced tokens for speed."""
    # Check cache first using hash of prompt + temperature
    prompt_hash = hash_text(f"{prompt}:{temperature}:{max_tokens}")
    cache_key = f"gemini_opt:{prompt_hash}"
    
    if cache_key in QUERY_CACHE:
        logger.info("Cache hit for optimized Gemini API call")
        return QUERY_CACHE[cache_key]
    
    wait_for_rate_limit_optimized()

    if not _ensure_model_initialized():
        logger.error("Gemini model unavailable; returning empty response")
        return ""

    # Reduced retry attempts to avoid excessive API calls
    max_retries = 2
    base_delay = 3.0  # Increased base delay

    inflight_evt = _wait_for_inflight(cache_key)
    last_key_index = _ACTIVE_GEMINI_INDEX
    try:
        for attempt in range(max_retries):
            try:
                if _llm_semaphore is not None:
                    _llm_semaphore.acquire()
                _log_gemini_api_call(
                    context="rate_limited_generate_content_optimized",
                    prompt=prompt,
                    temperature=temperature,
                    max_tokens=max_tokens,
                    optimized=True,
                )
                response = client.models.generate_content(
                    model=MODEL_ID,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        thinking_config=types.ThinkingConfig(thinking_level="low"),
                        temperature=temperature,
                        max_output_tokens=max_tokens,
                    )
                )
                if response and response.text:
                    result = response.text.strip()
                    QUERY_CACHE[cache_key] = result
                    if len(QUERY_CACHE) % 20 == 0:
                        save_query_cache()
                    return result
                else:
                    logger.warning("Empty response from Gemini API")
                    if attempt < max_retries - 1:
                        time.sleep(base_delay * (2 ** attempt))  # Exponential backoff
                    continue
            except Exception as e:
                error_str = str(e).lower()
                error_message = str(e)
                is_rate_limit = any(k in error_str for k in ["429", "resourceexhausted", "quota", "rate limit", "rate_limit"])
                is_quota_exhausted = _is_quota_exhausted(error_message)
                
                if attempt == max_retries - 1:
                    logger.error(f"All retries exhausted. Last error: {e}")
                    break
                
                # Extract retry delay from API response
                retry_delay = _extract_retry_delay(error_message)
                
                if is_rate_limit or is_quota_exhausted:
                    # Mark current key as exhausted if quota exhausted
                    if is_quota_exhausted and _ACTIVE_GEMINI_INDEX is not None:
                        _mark_key_exhausted(_ACTIVE_GEMINI_INDEX)
                        logger.error(f"Key #{_ACTIVE_GEMINI_INDEX + 1} quota exhausted. Marked for cooldown.")
                    
                    # Use API-provided retry delay if available, otherwise use exponential backoff
                    if retry_delay:
                        backoff_time = min(retry_delay, 120)  # Cap at 2 minutes
                        logger.warning(f"Quota/Rate limit detected on key #{_ACTIVE_GEMINI_INDEX + 1 if _ACTIVE_GEMINI_INDEX is not None else 'unknown'}. API suggests retry in {retry_delay}s. Waiting {backoff_time}s...")
                    else:
                        backoff_time = min(base_delay * (2 ** attempt), 60)  # Max 60s for exponential backoff
                        logger.warning(f"Rate limit detected on key #{_ACTIVE_GEMINI_INDEX + 1 if _ACTIVE_GEMINI_INDEX is not None else 'unknown'} (attempt {attempt+1}/{max_retries}), waiting {backoff_time:.1f}s")
                    
                    # Try switching to a different key (even for quota exhaustion, as keys might be from different accounts)
                    if _try_all_keys_rotation() and _ACTIVE_GEMINI_INDEX != last_key_index:
                        logger.info(f"Rotated to key #{_ACTIVE_GEMINI_INDEX + 1} after quota/rate limit. Retrying immediately...")
                        last_key_index = _ACTIVE_GEMINI_INDEX
                        # Don't wait if we successfully rotated to a new key
                        continue
                    else:
                        # All keys exhausted or no other keys available
                        if is_quota_exhausted:
                            logger.error(f"All keys appear exhausted. Waiting {backoff_time}s before retry...")
                        else:
                            logger.warning(f"All keys rate limited. Waiting {backoff_time}s before retry...")
                        time.sleep(backoff_time)
                        continue
                else:
                    # Non-rate-limit error
                    logger.error(f"API error: {e}")
                    if _switch_to_fallback_key() and _ACTIVE_GEMINI_INDEX != last_key_index:
                        logger.info(f"Switched to fallback Gemini API key #{_ACTIVE_GEMINI_INDEX + 1} after error")
                        last_key_index = _ACTIVE_GEMINI_INDEX
                        time.sleep(2)  # Short delay before retry
                        continue
                    else:
                        time.sleep(base_delay * (2 ** attempt))  # Exponential backoff
                        continue
            finally:
                try:
                    if _llm_semaphore is not None:
                        _llm_semaphore.release()
                except Exception:
                    pass
    finally:
        if inflight_evt is not None and _inflight_lock is not None:
            with _inflight_lock:
                try:
                    evt = _inflight_prompts.pop(cache_key, None)
                    if evt is not None:
                        evt.set()
                except Exception:
                    pass
    
    # All retries failed - return empty response
    logger.error("All Gemini API attempts failed. Quota may be exhausted.")
    return ""

@timing_decorator
def get_embedding(text: str) -> np.ndarray:
    """Get embeddings using sentence-transformers model."""
    if len(text.split()) > 8000:
        text = " ".join(text.split()[:8000])
    try:
        embeddings = embedding_model.encode([text], convert_to_numpy=True)
        return embeddings[0]
    except Exception as e:
        logger.error(f"Error generating embedding: {e}")
        return None

def save_embeddings(embeddings: List[np.ndarray], document_map: List[Dict[str, Any]]) -> None:
    """Save embeddings and document mapping to disk."""
    try:
        embeddings_array = np.array(embeddings)
        np.save(EMBEDDING_CACHE_FILE, embeddings_array)
        with open(DOCUMENT_MAP_FILE, 'w') as f:
            json.dump(document_map, f, indent=2)
        logger.info(f"Saved {len(embeddings)} embeddings to {EMBEDDING_CACHE_FILE}")
    except Exception as e:
        logger.error(f"Error saving embeddings: {e}")

def load_embeddings() -> Tuple[Optional[np.ndarray], Optional[List[Dict[str, Any]]]]:
    """Load embeddings and document mapping from disk."""
    try:
        if os.path.exists(EMBEDDING_CACHE_FILE) and os.path.exists(DOCUMENT_MAP_FILE):
            embeddings = np.load(EMBEDDING_CACHE_FILE)
            with open(DOCUMENT_MAP_FILE, 'r') as f:
                document_map = json.load(f)
            logger.info(f"Loaded {len(embeddings)} embeddings from cache")
            return embeddings, document_map
    except Exception as e:
        logger.error(f"Error loading embeddings: {e}")
    return None, None

def build_and_save_faiss_index(embeddings: np.ndarray) -> Optional[faiss.Index]:
    """Build and save optimized FAISS index for fast retrieval."""
    try:
        dimension = embeddings.shape[1]
        n_vectors = embeddings.shape[0]
        
        logger.info(f"Building optimized FAISS index for {n_vectors} vectors, dimension {dimension}")
        
        # For small datasets (< 10k), use flat index for speed
        if n_vectors < 10000:
            index = faiss.IndexFlatL2(dimension)
            index.add(embeddings)
            logger.info("Using IndexFlatL2 for optimal small dataset performance")
        else:
            # For larger datasets, use optimized IVF
            n_clusters = min(int(np.sqrt(n_vectors) * 2), n_vectors // 20)
            n_clusters = max(n_clusters, 4)
            
            # Create optimized quantizer
            quantizer = faiss.IndexFlatL2(dimension)
            index = faiss.IndexIVFFlat(quantizer, dimension, n_clusters)
            
            # Optimize search parameters for speed
            index.nprobe = max(1, n_clusters // 8)  # Faster search with slight accuracy trade-off
            
            # Train and add vectors
            logger.info(f"Training IVF index with {n_clusters} clusters, nprobe={index.nprobe}")
            index.train(embeddings)
            index.add(embeddings)
        
        # Enable multi-threading for search
        faiss.omp_set_num_threads(4)
        
        faiss.write_index(index, FAISS_INDEX_FILE)
        logger.info(f"Saved optimized FAISS index with {index.ntotal} vectors")
        return index
        
    except Exception as e:
        logger.error(f"Error building optimized FAISS index: {e}")
        return None

def load_faiss_index() -> Optional[faiss.Index]:
    """Load FAISS index from disk."""
    try:
        if os.path.exists(FAISS_INDEX_FILE):
            index = faiss.read_index(FAISS_INDEX_FILE)
            logger.info(f"Loaded FAISS index with {index.ntotal} vectors")
            return index
    except Exception as e:
        logger.error(f"Error loading FAISS index: {e}")
    return None

def process_documents(new_document_path: Optional[str] = None) -> Tuple[List[str], np.ndarray, Any]:
    """Process documents and generate embeddings with optimized segmentation and caching.
    If new_document_path is provided, only process that document and merge with existing embeddings."""
    
    # Ensure required directories exist
    for directory in [EMBEDDING_DIR, INDEX_DIR, CACHE_DIR, PDF_FOLDER]:
        if not os.path.exists(directory):
            os.makedirs(directory)
            logger.info(f"Created directory: {directory}")
    
    logger.info("Checking for existing embeddings...")
    cached_embeddings, cached_doc_map = load_embeddings()
    
    all_segments = []
    all_embeddings = []
    document_map = []
    
    # If we have cached embeddings, use them as a base
    if cached_embeddings is not None and cached_doc_map is not None:
        logger.info(f"Loaded {len(cached_embeddings)} embeddings from cache")
        # Extract segments from document map (not full text!)
        all_segments = []
        for doc in cached_doc_map:
            if "segments" in doc and isinstance(doc["segments"], list):
                all_segments.extend(doc["segments"])
            else:
                # Fallback: if segments not in map, use full text as single segment
                logger.warning(f"Document '{doc.get('filename', 'unknown')}' has no segments, using full text")
                all_segments.append(doc["text"])
        all_embeddings = list(cached_embeddings)
        document_map = cached_doc_map.copy()
        logger.info(f"Found existing embeddings! Extracted {len(all_segments)} segments from {len(cached_doc_map)} documents")
    
    # If a new document path is provided, only process that document
    if new_document_path:
        logger.info(f"Processing new document: {new_document_path}")
        try:
            # Extract text from the document
            if new_document_path.lower().endswith('.pdf'):
                text = extract_text_from_pdf(new_document_path)
            elif new_document_path.lower().endswith('.docx'):
                text = extract_text_from_docx(new_document_path)
            else:
                logger.error(f"Unsupported file format: {new_document_path}")
                return None, None, None

            if not text:
                logger.error(f"No text content could be extracted from {new_document_path}")
                return None, None, None

            # Process the new document
            segments = process_text_into_segments(text)
            logger.info(f"Created {len(segments)} segments from {new_document_path}")

            # Generate embeddings for new segments
            batch_size = 32
            for i in range(0, len(segments), batch_size):
                batch = segments[i:i+batch_size]
                try:
                    batch_embeddings = embedding_model.encode(batch, convert_to_numpy=True)
                    all_embeddings.extend(batch_embeddings)
                    all_segments.extend(batch)
                except Exception as e:
                    logger.error(f"Error generating batch embeddings: {e}")
                    continue
                logger.info(f"Processed batch {i//batch_size + 1}/{(len(segments)-1)//batch_size + 1}")

            # Add new document to document map
            document_map.append({
                "filename": os.path.basename(new_document_path),
                "text": text,
                "segments": segments,
                "position": len(all_segments) - len(segments)
            })

        except Exception as e:
            logger.error(f"Error processing new document: {e}")
            return None, None, None
    else:
        # Check if we already have embeddings and they cover all current documents
        frameworks_dir = Path(__file__).parent / "compliance_frameworks"
        if not frameworks_dir.exists():
            logger.error(f"compliance_frameworks directory not found at: {frameworks_dir.absolute()}")
            return None, None, None

        # Get list of all PDF and DOCX files
        framework_files = list(frameworks_dir.glob("*.pdf")) + list(frameworks_dir.glob("*.docx"))
        current_filenames = {file_path.name for file_path in framework_files}
        
        # If we have cached embeddings, check if they cover all current files
        if cached_embeddings is not None and cached_doc_map is not None:
            cached_filenames = {doc["filename"] for doc in cached_doc_map}
            
            # If all current files are already cached, return cached data
            if current_filenames.issubset(cached_filenames):
                logger.info(f"All {len(current_filenames)} documents are already processed. Using cached data.")
                # Load and return the existing FAISS index
                index = load_faiss_index()
                if index is not None:
                    logger.info(f"Loaded FAISS index with {index.ntotal} vectors")
                    return all_segments, cached_embeddings, index
                else:
                    logger.info("FAISS index not found, rebuilding...")
                    index = build_and_save_faiss_index(cached_embeddings)
                    return all_segments, cached_embeddings, index
            else:
                # Some files are new, process only the new ones
                new_files = current_filenames - cached_filenames
                logger.info(f"Found {len(new_files)} new documents to process: {new_files}")
                
                # Process only new files
                for file_path in framework_files:
                    if file_path.name in new_files:
                        try:
                            logger.info(f"Processing {file_path.name}")
                            if file_path.suffix.lower() == '.pdf':
                                text = extract_text_from_pdf(str(file_path))
                            else:  # .docx
                                text = extract_text_from_docx(str(file_path))

                            if not text:
                                logger.warning(f"No text content extracted from {file_path.name}")
                                continue

                            segments = process_text_into_segments(text)
                            logger.info(f"Created {len(segments)} segments from {file_path.name}")

                            # Generate embeddings in batches
                            batch_size = 32
                            for i in range(0, len(segments), batch_size):
                                batch = segments[i:i+batch_size]
                                try:
                                    batch_embeddings = embedding_model.encode(batch, convert_to_numpy=True)
                                    all_embeddings.extend(batch_embeddings)
                                    all_segments.extend(batch)
                                except Exception as e:
                                    logger.error(f"Error generating batch embeddings: {e}")
                                    continue
                                logger.info(f"Processed batch {i//batch_size + 1}/{(len(segments)-1)//batch_size + 1}")

                            # Add to document map
                            document_map.append({
                                "filename": file_path.name,
                                "text": text,
                                "segments": segments,
                                "position": len(all_segments) - len(segments)
                            })

                        except Exception as e:
                            logger.error(f"Error processing {file_path.name}: {e}")
                            continue
        else:
            # No cached embeddings, process all documents
            logger.info("No cached embeddings found. Processing all documents...")
            for file_path in framework_files:
                try:
                    logger.info(f"Processing {file_path.name}")
                    if file_path.suffix.lower() == '.pdf':
                        text = extract_text_from_pdf(str(file_path))
                    else:  # .docx
                        text = extract_text_from_docx(str(file_path))

                    if not text:
                        logger.warning(f"No text content extracted from {file_path.name}")
                        continue

                    segments = process_text_into_segments(text)
                    logger.info(f"Created {len(segments)} segments from {file_path.name}")

                    # Generate embeddings in batches
                    batch_size = 32
                    for i in range(0, len(segments), batch_size):
                        batch = segments[i:i+batch_size]
                        try:
                            batch_embeddings = embedding_model.encode(batch, convert_to_numpy=True)
                            all_embeddings.extend(batch_embeddings)
                            all_segments.extend(batch)
                        except Exception as e:
                            logger.error(f"Error generating batch embeddings: {e}")
                            continue
                        logger.info(f"Processed batch {i//batch_size + 1}/{(len(segments)-1)//batch_size + 1}")

                    # Add to document map
                    document_map.append({
                        "filename": file_path.name,
                        "text": text,
                        "segments": segments,
                        "position": len(all_segments) - len(segments)
                    })

                except Exception as e:
                    logger.error(f"Error processing {file_path.name}: {e}")
                    continue

    if not all_segments:
        logger.info("No segments were created. Please check your documents.")
        return None, None, None
    
    # Only save and rebuild index if we processed new documents
    if new_document_path or (cached_embeddings is None) or len(document_map) > len(cached_doc_map or []):
        logger.info(f"Updating embeddings and index with {len(all_embeddings)} total embeddings")
        embeddings_array = np.array(all_embeddings)
        
        # Save the updated embeddings and document map
        save_embeddings(all_embeddings, document_map)
        
        # Build and save the new index
        index = build_and_save_faiss_index(embeddings_array)
        
        logger.info(f"Returning {len(all_segments)} total segments")
        return all_segments, embeddings_array, index
    else:
        # Return cached data with existing index
        index = load_faiss_index()
        if index is None:
            logger.info("Rebuilding FAISS index from cached embeddings")
            embeddings_array = np.array(all_embeddings)
            index = build_and_save_faiss_index(embeddings_array)
        
        return all_segments, np.array(all_embeddings), index

def detect_concise_request(query: str) -> bool:
    """Detect if user is asking for a concise/brief response."""
    concise_indicators = [
        "concisely", "briefly", "short", "quick", "summarize", "summary",
        "in brief", "quick answer", "short answer", "bullet points",
        "key points", "main points", "overview", "tldr", "tl;dr"
    ]
    query_lower = query.lower()
    return any(indicator in query_lower for indicator in concise_indicators)

def get_concise_max_tokens(query: str) -> int:
    """Return appropriate max_tokens based on whether user wants concise response."""
    if detect_concise_request(query):
        return 512  # Much shorter for concise requests
    return 4000  # Increased for comprehensive lifecycle and detailed responses

@timing_decorator
def expert_security_controls(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for security controls with evidence-based prompting."""
    print("\n" + "="*80)
    print("🔒 SECURITY EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"🔒 SECURITY EXPERT triggered for query: {query[:100]}")
    
    is_concise = detect_concise_request(query)
    max_tokens = get_concise_max_tokens(query)
    
    if is_concise:
        prompt = (
            "Provide a CONCISE, bullet-point response for this security/compliance query.\n\n"
            f"Query: {query}\n"
            f"Context: {context[:500]}\n\n"
            "Format as:\n"
            "**Key Steps:**\n"
            "• Point 1 (cite source if from context)\n"
            "• Point 2 (cite source if from context)\n"
            "• Point 3 (cite source if from context)\n\n"
            "Keep it under 150 words total. Focus on actionable steps only."
        )
    else:
        prompt = (
            "You are a cybersecurity and information security expert specializing in enterprise security controls and compliance frameworks.\n\n"
            f"Previous conversation context:\n{conversation_context}\n\n"
            f"Current Query: {query}\n\n"
            f"SECURITY FRAMEWORK DOCUMENTS:\n{context}\n\n"
            "CRITICAL INSTRUCTIONS - HYBRID EVIDENCE-BASED RESPONSE:\n\n"
            "**FOLLOW-UP QUESTION HANDLING:**\n"
            "If the current query is a follow-up question (references previous conversation, uses pronouns like 'it', 'that', 'those', 'this', or asks for clarification/expansion):\n"
            "1. FIRST, review the previous conversation context above to understand what was discussed\n"
            "2. Reference the previous answer when responding (e.g., 'As mentioned earlier...', 'Building on the previous discussion about...')\n"
            "3. If the user asks about something from a previous answer, provide additional details or clarification\n"
            "4. Maintain continuity with the previous conversation while answering the new question\n"
            "5. If the follow-up is about a specific framework/control mentioned earlier, focus on that context\n\n"
            "You MUST provide a COMPLETE answer. Structure your response as follows:\n\n"
            "**Framework Requirements**\n"
            "1. If relevant information exists in the framework documents, quote the EXACT text using format: [Your statement] (Evidence: <span style=\"color:#008000\">\"exact quote\"</span> - <span style=\"color:#008000\">Framework Name/Control ID</span>)\n"
            "2. ALWAYS highlight document-derived evidence AND control IDs/article numbers in green using <span style=\"color:#008000\">text</span>\n"
            "3. If no relevant information exists in the documents, skip this section entirely and proceed with the next section.\n\n"
            "**Implementation Guidance**\n"
            "1. Provide complete information about the query topic using your expert knowledge.\n"
            "2. Provide comprehensive details, implementation guidance, and best practices.\n"
            "3. When mentioning control IDs, article numbers, or framework requirements, ALWAYS highlight them in green: <span style=\"color:#008000\">ISO 27001 A.9.2.1</span> or <span style=\"color:#008000\">Article 17 GDPR</span>\n"
            "4. Use this format: [Information] (Based on <span style=\"color:#008000\">{Framework} Control/Article</span> standards and industry best practices)\n\n"
            "**RESPONSE STRUCTURE:**\n"
            "1. Start with a brief overview\n"
            "2. **Framework Requirements:** (Cite exact quotes from context if available, otherwise skip)\n"
            "3. **Implementation Guidance:** (Use your knowledge to provide complete answer with practical steps)\n\n"
            "EXAMPLE FORMAT:\n"
            "**Framework Requirements:**\n"
            "ISO 27001 requires access control policies (Evidence: <span style=\"color:#008000\">\"Access control policy\"</span> - <span style=\"color:#008000\">ISO 27001 A.9.1.1</span>)\n\n"
            "**Implementation Guidance:**\n"
            "Control <span style=\"color:#008000\">A.9.2.1</span> specifically covers user access management requirements.\n\n"
            "**Implementation Guidance (Based on ISO 27001 Standards):**\n"
            "Control A.9 specifically covers Access Control and includes:\n"
            "- A.9.1: Business requirements of access control\n"
            "- A.9.2: User access management\n"
            "- A.9.3: User responsibilities\n"
            "- A.9.4: System and application access control\n"
            "Organizations need to implement user access management processes including user registration, privilege management, and access review procedures. (Based on ISO 27001 standards and industry best practices)\n\n"
            "Provide a comprehensive analysis covering:\n"
            "1. **Security Controls & Requirements**: Specific controls from NIST, ISO 27001, CIS (cite documents first, then supplement)\n"
            "2. **Implementation Guidelines**: Step-by-step technical implementation (cite framework requirements, add practical steps)\n"
            "3. **Risk Assessment**: Threats, vulnerabilities, and risk levels (reference documents, add expert insights)\n"
            "4. **Monitoring & Validation**: Methods to verify control effectiveness (cite verification procedures, add best practices)\n"
            "5. **Best Practices**: Industry-proven measures (distinguish framework requirements vs. best practices)\n"
            "6. **Compliance Mapping**: How controls map to regulations (cite specific control mappings, add context)\n\n"
            "IMPORTANT RULES:\n"
            "- DO NOT say 'not in documents', 'not available', 'there is no direct mention', or explain what's missing. Simply provide the information.\n"
            "- DO NOT explain your approach or methodology. Just provide the answer directly.\n"
            "- Always provide a COMPLETE answer. Never leave the user with incomplete information.\n"
            "- If documents have partial information, cite it first, then supplement with your knowledge to provide a full answer.\n\n"
            "Provide a direct, complete answer without meta-commentary:"
        )
    return rate_limited_generate_content_optimized(prompt, max_tokens=max_tokens)

@timing_decorator
def expert_privacy_regulations(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for privacy regulations with evidence-based prompting."""
    print("\n" + "="*80)
    print("⚖️ PRIVACY EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"⚖️ PRIVACY EXPERT triggered for query: {query[:100]}")
    
    prompt = (
        "You are a data privacy and protection expert with deep knowledge of GDPR, CCPA, and global data governance.\n\n"
        f"Previous conversation context:\n{conversation_context}\n\n"
        f"Current Query: {query}\n\n"
        f"PRIVACY REGULATION DOCUMENTS:\n{context}\n\n"
        "CRITICAL INSTRUCTIONS - OPERATIONALIZING PRIVACY:\n"
        "1. **Lifecycle Alignment:** If the user asks for a data lifecycle (Collection -> Deletion), structure your answer exactly in that order.\n"
        "2. **Tech-Specific Implementation:** When discussing privacy controls (e.g., Right to Erasure), provide the technical method relevant to the user's stack found in context (e.g., 'Use AWS S3 Lifecycle Policies for automated deletion').\n"
        "3. **Ignore Irrelevant Stack Info:** Do not recommend tools from cloud providers the user is not using unless explicitly comparing them.\n\n"
        "You MUST provide a COMPLETE answer. Structure your response as follows:\n\n"
            "**Regulatory Requirements**\n"
            "1. If relevant legal text exists in the regulation documents, quote it using format: [Legal requirement] (Legal Basis: <span style=\"color:#008000\">\"exact quote\"</span> - <span style=\"color:#008000\">Regulation Article/Section</span>).\n"
            "2. ALWAYS highlight article numbers and section references in green: <span style=\"color:#008000\">Article 17 GDPR</span> or <span style=\"color:#008000\">CCPA Section 1798.105</span>\n"
            "3. CRITICAL: If no relevant information exists in the documents, DO NOT include this section at all. Do not write 'There is no direct mention', 'not in documents', or any explanation about missing information. Simply skip this section and go directly to the next section.\n\n"
            "**Technical Implementation**\n"
            "1. Provide complete information about the query topic using your expert knowledge.\n"
            "2. Translate legal requirements into 'Technical Action' matching the user's cloud provider.\n"
            "3. Provide comprehensive details on how to execute the right/obligation technically.\n"
            "4. When mentioning article numbers or regulation sections, ALWAYS highlight them in green: <span style=\"color:#008000\">Article 17</span>, <span style=\"color:#008000\">Article 32</span>\n"
            "5. Use format: [Technical Action] (Based on <span style=\"color:#008000\">{Regulation} Article/Control</span> standards and {Cloud Provider} best practices).\n\n"
        "**RESPONSE STRUCTURE:**\n"
        "1. **Regulatory Requirements**: Cite Article numbers and legal text (highlight quotes in green span). SKIP THIS SECTION ENTIRELY if no relevant information exists in the documents - do not mention its absence.\n"
        "2. **Technical Implementation**: Provide privacy configurations and technical actions for the user's tech stack.\n"
        "3. **Data Subject Rights**: Explain how to handle specific requests (Access, Deletion) with evidence citations.\n"
        "4. **Documentation**: Outline RoPA or record-keeping requirements.\n\n"
            "ADDITIONAL RULES:\n"
            "- CRITICAL: If the 'Regulatory Requirements' section has no content from documents, DO NOT include that section header at all. Do not write 'There is no direct mention', 'not in documents', 'not available', or any explanation about missing information. Simply skip that section entirely and start directly with 'Technical Implementation'.\n"
            "- DO NOT say 'not in documents', 'not available', 'there is no direct mention', or explain what's missing. Simply provide the information or skip the section.\n"
            "- DO NOT explain your approach or methodology. Just provide the answer directly.\n"
            "- Maintain continuity with conversation context for follow-up questions.\n"
            "- ALWAYS highlight ALL article numbers, control IDs, and framework references in green: <span style=\"color:#008000\">Article 17 GDPR</span>, <span style=\"color:#008000\">Article 32</span>, <span style=\"color:#008000\">CCPA Section 1798.105</span>\n"
            "- Highlight document-derived evidence quotes using <span style=\"color:#008000\">\"quote\"</span> and include Article references in green.\n"
            "- Example: Data subjects have the right to erasure (\"right to be forgotten\") (<span style=\"color:#008000\">Article 17 GDPR</span>).\n"
            "- Ensure each legal obligation is paired with a concrete technical execution path for the specified stack.\n\n"
            "Provide a direct, complete answer without meta-commentary. If no document information exists, skip the 'Regulatory Requirements' section entirely:"
    )
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def expert_audit_compliance(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for audit compliance with evidence-based prompting."""
    print("\n" + "="*80)
    print("📄 AUDIT EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"📄 AUDIT EXPERT triggered for query: {query[:100]}")
    
    # Check if we have substantial framework context
    has_framework_docs = context and len(context.strip()) > 200
    
    if has_framework_docs:
        # When we have framework documents, cite them
        prompt = (
            "You are an audit and compliance expert with deep expertise in enterprise risk management, ISO 27001, SOC 2, and cloud security.\n\n"
            f"Previous conversation context:\n{conversation_context}\n\n"
            f"Current Query: {query}\n\n"
            f"COMPLIANCE FRAMEWORK DOCUMENTS:\n{context}\n\n"
            "CRITICAL INSTRUCTIONS:\n"
            "1. **Tech Stack Validation:** Check the 'Current Query' and 'Previous conversation context' for the user's specific technology stack (e.g., AWS, Azure, GCP, On-Prem).\n"
            "2. **Filter Retrieved Docs:** IF the user specifies a technology (e.g., AWS), you must IGNORE or contextualize information from 'COMPLIANCE FRAMEWORK DOCUMENTS' that refers to competing technologies unless explicitly comparing them.\n"
            "3. **Specific Control Mapping:** When referencing ISO 27001 or other frameworks, cite the specific alphanumeric Control ID (e.g., 'Control A.12.3.1', 'Article 32').\n\n"
            "You MUST provide a COMPLETE answer. Structure your response as follows:\n\n"
            "**Control Requirements**\n"
            "1. If relevant information exists in the documents, provide a NATURAL, PARAPHRASED statement explaining the requirement, then cite the exact quote as evidence.\n"
            "2. Format: [Natural paraphrased statement explaining the requirement in your own words] (Evidence: <span style=\"color:#008000\">\"exact quote from document\"</span> - <span style=\"color:#008000\">Control ID A.X.Y</span>).\n"
            "3. CRITICAL: The statement before the evidence MUST be a natural explanation or summary, NOT an exact copy of the evidence quote. Paraphrase the requirement in clear, natural language.\n"
            "4. Example of CORRECT format:\n"
            "   Organizations must include cloud providers in their supplier security policies to manage risks from third-party data access (Evidence: <span style=\"color:#008000\">\"The cloud service customer should include the cloud service provider as a type of supplier in its information security policy for supplier relationships.\"</span> - <span style=\"color:#008000\">ISO 27001 A.15.1.2</span>).\n"
            "5. Example of INCORRECT format (DO NOT USE):\n"
            "   The cloud service customer should include the cloud service provider as a type of supplier in its information security policy for supplier relationships (Evidence: <span style=\"color:#008000\">\"The cloud service customer should include the cloud service provider as a type of supplier in its information security policy for supplier relationships.\"</span>).\n"
            "6. ALWAYS highlight control IDs in green: <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>, <span style=\"color:#008000\">A.12.3.1</span>\n"
            "7. If no relevant information exists in the documents, skip this section entirely and proceed with the next section.\n\n"
            "**Implementation Details**\n"
            "1. Provide complete information about the query topic using your expert knowledge.\n"
            "2. Map the compliance requirement to the user's specific tools.\n"
            "3. Example: If the user is on AWS and the requirement is 'Encryption', recommend 'AWS KMS' and 'S3 Server-Side Encryption'.\n"
            "4. When mentioning control IDs or article numbers, ALWAYS highlight them in green: <span style=\"color:#008000\">ISO 27001 Control A.9.2.1</span>, <span style=\"color:#008000\">Article 32</span>\n"
            "5. Use format: [Technical Implementation] (Standard: <span style=\"color:#008000\">ISO 27001 Control A.X.Y</span>, Best Practice for [User's Tech Stack]).\n\n"
            "**RESPONSE STRUCTURE:**\n"
            "1. **Audit Scope**: Identify the specific controls relevant to the query.\n"
            "2. **Control Requirements**: Provide natural, paraphrased statements explaining the requirements, then cite exact quotes as evidence with Control IDs (use green-highlight format). The statement should NOT be an exact copy of the evidence quote - paraphrase it naturally. Otherwise, provide requirements from your knowledge.\n"
            "3. **Technical Implementation**: Map controls to the user's specific cloud/tech stack.\n"
            "4. **Evidence for Auditors**: List the specific logs/artifacts the user must produce.\n"
            "   - CRITICAL: Format this section as a numbered list with each item on a separate line.\n"
            "   - Use HTML ordered list tags: <ol><li>Item 1</li><li>Item 2</li></ol> OR use numbered format with line breaks.\n"
            "   - Each evidence item must be on its own line, not in paragraph form.\n"
            "   - Example format:\n"
            "     1. Backup Policy: A documented backup policy that defines...\n"
            "     2. Backup Configuration: Configuration details of the backup systems...\n"
            "     3. Backup Logs: Logs showing that backups are being performed...\n\n"
            "ADDITIONAL RULES:\n"
            "- DO NOT say 'there is no direct mention', 'the document focuses on', 'not in documents', or explain what's missing. Simply provide the information.\n"
            "- DO NOT explain your approach or methodology. Just provide the answer directly.\n"
            "- Never leave the user with generic advice. If they use AWS, provide AWS-specific audit procedures.\n"
            "- ALWAYS highlight ALL control IDs, article numbers, and framework references in green: <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>, <span style=\"color:#008000\">Article 32</span>, <span style=\"color:#008000\">SOC 2 CC6.1</span>\n"
            "- Highlight document-derived evidence quotes using <span style=\"color:#008000\">\"quote\"</span> and include the control ID/article number in green.\n"
            "- Example: Organizations must implement access controls (<span style=\"color:#008000\">ISO 27001 A.9.2.1</span>).\n"
            "- Maintain continuity for follow-up questions by referencing the conversation context when applicable.\n\n"
            "Provide a direct, complete answer without meta-commentary:"
        )
    else:
        # When no framework documents are provided, answer based on general knowledge
        prompt = (
            "You are an audit and compliance expert with expertise in enterprise risk management and regulatory compliance frameworks.\n\n"
            f"Previous conversation context:\n{conversation_context}\n\n"
            f"Current Query: {query}\n\n"
            "Provide a comprehensive, professional analysis of the requested compliance framework or topic.\n\n"
            "Structure your response with:\n"
            "1. **Overview**: Brief introduction to the framework or compliance area\n"
            "2. **Key Requirements**: Main compliance requirements and standards\n"
            "3. **Audit Requirements**: Specific audit standards and procedures\n"
            "4. **Evidence Collection**: Documentation and artifacts typically needed\n"
            "5. **Compliance Verification**: Common methods to assess and validate compliance\n"
            "6. **Risk Assessment Framework**: How to identify, analyze, and prioritize risks\n"
            "7. **Control Testing**: Methods to test and validate control effectiveness\n"
            "8. **Remediation Planning**: Best practices for addressing gaps and findings\n"
            "9. **Continuous Monitoring**: Strategies for ongoing compliance maintenance\n\n"
            "Format your response with clear headings (using ##) and bullet points. Be specific, actionable, and professional.\n"
            "Provide comprehensive information based on established compliance standards and best practices.\n"
        )
    
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def expert_financial_compliance(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for financial compliance with evidence-based prompting and chain-of-thought."""
    print("\n" + "="*80)
    print("💰 FINANCIAL EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"💰 FINANCIAL EXPERT triggered for query: {query[:100]}")
    
    prompt = (
        "You are a financial compliance expert with expertise in banking regulations, payment standards, and financial reporting requirements.\n\n"
        f"Previous conversation context:\n{conversation_context}\n\n"
        f"Current Query: {query}\n\n"
        f"FINANCIAL REGULATION DOCUMENTS:\n{context}\n\n"
        "CRITICAL INSTRUCTIONS - HYBRID EVIDENCE-BASED RESPONSE:\n\n"
        "**FOLLOW-UP QUESTION HANDLING:**\n"
        "If the current query is a follow-up question (references previous conversation, uses pronouns like 'it', 'that', 'those', 'this', or asks for clarification/expansion):\n"
        "1. FIRST, review the previous conversation context above to understand what was discussed\n"
        "2. Reference the previous answer when responding (e.g., 'As mentioned earlier...', 'Building on the previous discussion about...')\n"
        "3. If the user asks about something from a previous answer, provide additional details or clarification\n"
        "4. Maintain continuity with the previous conversation while answering the new question\n"
        "5. If the follow-up is about a specific regulation/requirement mentioned earlier, focus on that context\n\n"
            "You MUST provide a COMPLETE answer. Structure your response as follows:\n\n"
            "**Regulatory Requirements**\n"
            "1. If relevant information exists in the regulation documents, quote the EXACT text using format: [Requirement] (Regulation: <span style=\"color:#008000\">\"exact quote\"</span> - <span style=\"color:#008000\">Standard/Section</span>)\n"
            "2. ALWAYS highlight document-derived evidence AND regulation sections/requirements in green using <span style=\"color:#008000\">text</span>\n"
            "3. If no relevant information exists in the documents, skip this section entirely and proceed with the next section.\n\n"
            "**Implementation Guidance**\n"
            "1. Provide complete information about the query topic using your expert knowledge.\n"
            "2. Provide comprehensive details, implementation guidance, and best practices.\n"
            "3. When mentioning regulation requirements, sections, or control IDs, ALWAYS highlight them in green: <span style=\"color:#008000\">PCI DSS Requirement 1.1</span>, <span style=\"color:#008000\">SOX Section 404</span>\n"
            "4. Use this format: [Information] (Based on <span style=\"color:#008000\">{Regulation} Requirement/Section</span> standards and industry best practices)\n\n"
        "**RESPONSE STRUCTURE:**\n"
        "1. Start with a brief overview\n"
        "2. **Information from Regulation Documents:** (Cite exact quotes from context)\n"
        "3. **Additional Information:** (Use your knowledge to complete the answer)\n"
        "4. **Implementation Guidance:** (Practical steps and recommendations)\n\n"
        "EXAMPLE FORMAT:\n"
        "**Information from Regulation Documents:**\n"
        "PCI DSS requires organizations to maintain a firewall configuration to protect cardholder data (Regulation: <span style=\"color:#008000\">\"Install and maintain a firewall configuration to protect cardholder data\"</span> - <span style=\"color:#008000\">PCI DSS Requirement 1.1</span>)\n\n"
        "**Additional Information Example:**\n"
        "Organizations must implement encryption for cardholder data (<span style=\"color:#008000\">PCI DSS Requirement 3.4</span>).\n\n"
        "**Additional Information (Based on PCI DSS Standards):**\n"
        "Requirement 1.1 includes establishing formal procedures for approving and testing network connections, documenting firewall rules, and reviewing firewall configurations at least every six months. Organizations must also restrict inbound and outbound traffic to only what is necessary for business purposes. (Based on PCI DSS standards and industry best practices)\n\n"
        "Focus on:\n"
        "1. **Financial Regulations**: PCI DSS, SOX, Basel III requirements (cite documents first, then supplement)\n"
        "2. **AML/KYC**: Anti-money laundering and Know Your Customer procedures (quote documents, add practical steps)\n"
        "3. **Payment Standards**: Card industry standards and requirements (cite documents, add implementation details)\n"
        "4. **Banking Compliance**: Financial services regulatory requirements (reference documents, add best practices)\n"
        "5. **Financial Reporting**: Disclosure and reporting requirements (cite documents, add reporting procedures)\n"
        "6. **Risk Management**: COSO, Basel frameworks (quote documents, add risk assessment methods)\n\n"
            "IMPORTANT RULES:\n"
            "- DO NOT say 'not in regulations', 'not available', 'there is no direct mention', or explain what's missing. Simply provide the information.\n"
            "- DO NOT explain your approach or methodology. Just provide the answer directly.\n"
            "- Always provide a COMPLETE answer. Never leave the user with incomplete information.\n"
            "- If documents have partial information, cite it first, then supplement with your knowledge to provide a full answer.\n\n"
            "Provide a direct, complete answer without meta-commentary:"
    )
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def expert_healthcare_compliance(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for healthcare compliance and HIPAA."""
    print("\n" + "="*80)
    print("🏥 HEALTHCARE EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"🏥 HEALTHCARE EXPERT triggered for query: {query[:100]}")
    
    is_concise = detect_concise_request(query)
    max_tokens = get_concise_max_tokens(query)
    has_framework_docs = context and len(context.strip()) > 200
    
    if is_concise:
        prompt = (
            "Provide a CONCISE, bullet-point response for HIPAA/healthcare compliance.\n\n"
            f"Query: {query}\n"
            f"Context: {context[:500] if has_framework_docs else 'General HIPAA knowledge'}\n\n"
            "Format as:\n"
            "**HIPAA Compliance Steps:**\n"
            "• Step 1\n"
            "• Step 2\n"
            "• Step 3\n\n"
            "Keep it under 150 words total. Focus on actionable steps only."
        )
    elif has_framework_docs:
        # When we have documents, cite them
        prompt = (
            "You are a healthcare compliance expert specializing in HIPAA and healthcare regulations.\n\n"
            f"Previous conversation context:\n{conversation_context}\n\n"
            f"Query: {query}\n\n"
            f"HEALTHCARE COMPLIANCE DOCUMENTS:\n{context}\n\n"
            "CRITICAL INSTRUCTIONS - HYBRID EVIDENCE-BASED RESPONSE:\n\n"
            "**FOLLOW-UP QUESTION HANDLING:**\n"
            "If the current query is a follow-up question (references previous conversation, uses pronouns like 'it', 'that', 'those', 'this', or asks for clarification/expansion):\n"
            "1. FIRST, review the previous conversation context above to understand what was discussed\n"
            "2. Reference the previous answer when responding (e.g., 'As mentioned earlier...', 'Building on the previous discussion about...')\n"
            "3. If the user asks about something from a previous answer, provide additional details or clarification\n"
            "4. Maintain continuity with the previous conversation while answering the new question\n"
            "5. If the follow-up is about a specific HIPAA rule/safeguard mentioned earlier, focus on that context\n\n"
            "You MUST provide a COMPLETE answer. Structure your response as follows:\n\n"
            "**Regulatory Requirements**\n"
            "1. If relevant information exists in the healthcare compliance documents, quote the EXACT text using format: [Requirement] (Evidence: <span style=\"color:#008000\">\"exact quote\"</span> - <span style=\"color:#008000\">HIPAA/Regulation Section</span>)\n"
            "2. ALWAYS highlight document-derived evidence AND section references in green using <span style=\"color:#008000\">text</span>\n"
            "3. If no relevant information exists in the documents, skip this section entirely and proceed with the next section.\n\n"
            "**Implementation Guidance**\n"
            "1. Provide complete information about the query topic using your expert knowledge.\n"
            "2. Provide comprehensive details, implementation guidance, and best practices.\n"
            "3. When mentioning HIPAA sections, safeguards, or regulation references, ALWAYS highlight them in green: <span style=\"color:#008000\">HIPAA §164.312(a)(1)</span>, <span style=\"color:#008000\">Technical Safeguard</span>\n"
            "4. Use this format: [Information] (Based on <span style=\"color:#008000\">HIPAA Section/Safeguard</span> standards and industry best practices)\n\n"
            "**RESPONSE STRUCTURE:**\n"
            "1. Start with a brief overview\n"
            "2. **Regulatory Requirements:** (Cite exact quotes from context if available, otherwise skip)\n"
            "3. **Implementation Guidance:** (Use your knowledge to provide complete answer with practical steps)\n\n"
            "Focus on:\n"
            "1. **HIPAA Privacy and Security Rules**: Cite documents first if available, then supplement with complete requirements\n"
            "2. **PHI Protection**: Data handling and security measures (cite documents if available, add technical safeguards)\n"
            "3. **Compliance Requirements**: Key obligations and standards (cite documents if available, add implementation details)\n"
            "4. **Risk Management**: Healthcare-specific risks and controls (cite documents if available, add risk assessment methods)\n"
            "5. **Best Practices**: Implementation guidance (cite documents if available, add practical steps)\n\n"
            "IMPORTANT RULES:\n"
            "- DO NOT say 'not in documents', 'not available', 'there is no direct mention', or explain what's missing. Simply provide the information.\n"
            "- DO NOT explain your approach or methodology. Just provide the answer directly.\n"
            "- Always provide a COMPLETE answer. Never leave the user with incomplete information.\n"
            "- If documents have partial information, cite it first, then supplement with your knowledge to provide a full answer.\n\n"
            "Provide a direct, complete answer without meta-commentary:"
        )
    else:
        # When no documents are provided, answer based on general knowledge
        prompt = (
            "You are a healthcare compliance expert specializing in HIPAA and healthcare regulations.\n\n"
            f"Previous conversation context:\n{conversation_context}\n\n"
            f"Query: {query}\n\n"
            "Provide a comprehensive, professional analysis of the healthcare compliance topic.\n\n"
            "Focus on:\n"
            "1. **Overview**: Introduction to the framework or compliance area\n"
            "2. **HIPAA Privacy and Security Rules**: Key requirements and standards\n"
            "3. **PHI Protection**: How to protect Protected Health Information\n"
            "4. **Covered Entities & Business Associates**: Obligations and requirements\n"
            "5. **Risk Assessment**: Healthcare-specific security risk management\n"
            "6. **Technical Safeguards**: Access controls, encryption, audit controls\n"
            "7. **Administrative Safeguards**: Policies, procedures, training\n"
            "8. **Physical Safeguards**: Facility access and device/media controls\n"
            "9. **Breach Notification**: Requirements and procedures\n"
            "10. **Compliance Verification**: Methods to assess and maintain compliance\n\n"
            "Format your response with clear headings (using ##) and bullet points. Be specific and actionable.\n"
            "Provide comprehensive information based on HIPAA standards and healthcare compliance best practices.\n"
        )
    return rate_limited_generate_content_optimized(prompt, max_tokens=max_tokens)

@timing_decorator
def expert_international_compliance(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for international and cross-border compliance."""
    print("\n" + "="*80)
    print("🌍 INTERNATIONAL EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"🌍 INTERNATIONAL EXPERT triggered for query: {query[:100]}")
    
    prompt = (
        "As an international compliance expert, analyze the following query:\n\n"
        f"{conversation_context}"
        f"Query: {query}\n"
        f"Context: {context}\n\n"
        "Focus on:\n"
        "1. Cross-border data transfer requirements\n"
        "2. International privacy laws (GDPR, LGPD, PIPEDA, etc.)\n"
        "3. Export control and trade compliance\n"
        "4. Multi-jurisdictional regulatory requirements\n"
        "5. Data localization and sovereignty issues\n"
        "6. International standards harmonization\n"
        "Chain-of-Thought Analysis:"
    )
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def expert_operational_compliance(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for operational compliance and business processes."""
    print("\n" + "="*80)
    print("⚙️ OPERATIONAL EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"⚙️ OPERATIONAL EXPERT triggered for query: {query[:100]}")
    
    prompt = (
        "As an operational compliance expert, analyze the following query:\n\n"
        f"{conversation_context}"
        f"Query: {query}\n"
        f"Context: {context}\n\n"
        "Focus on:\n"
        "1. Business process compliance and controls\n"
        "2. Vendor and third-party risk management\n"
        "3. Operational risk assessment and mitigation\n"
        "4. Business continuity and disaster recovery\n"
        "5. Change management and configuration control\n"
        "6. Incident response and breach notification\n"
        "Chain-of-Thought Analysis:"
    )
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def expert_industry_specific(query: str, context: str, conversation_context: str = "") -> str:
    """Expert analysis for industry-specific compliance requirements."""
    print("\n" + "="*80)
    print("🏭 INDUSTRY-SPECIFIC EXPERT TRIGGERED")
    print(f"Query: {query[:100]}...")
    print("="*80 + "\n")
    logger.info(f"🏭 INDUSTRY-SPECIFIC EXPERT triggered for query: {query[:100]}")
    
    prompt = (
        "As an industry-specific compliance expert, analyze the following query:\n\n"
        f"{conversation_context}"
        f"Query: {query}\n"
        f"Context: {context}\n\n"
        "Focus on:\n"
        "1. Industry-specific regulations (FERPA for education, GLBA for finance, etc.)\n"
        "2. Sector-specific standards and frameworks\n"
        "3. Professional licensing and certification requirements\n"
        "4. Industry best practices and guidance\n"
        "5. Regulatory body requirements and guidance\n"
        "6. Industry-specific risk factors and controls\n"
        "Think step by step and provide a detailed answer."
        
    )
    return rate_limited_generate_content_optimized(prompt)

@timing_decorator
def aggregate_expert_outputs(outputs: List[str], query: str, context: str, conversation_context: str = "") -> str:
    """Advanced aggregation and synthesis of expert outputs with cross-domain insights."""
    if not outputs:
        return "No expert analysis available."
    
    if len(outputs) == 1:
        return outputs[0]
    
    is_concise = detect_concise_request(query)
    max_tokens = get_concise_max_tokens(query)
    
    # Detect if this is a simple informational query (what is X, explain X, define X)
    query_lower = query.lower().strip()
    is_simple_info_query = any([
        query_lower.startswith('what is '),
        query_lower.startswith('what are '),
        query_lower.startswith('explain '),
        query_lower.startswith('define '),
        query_lower.startswith('describe '),
        'tell me about' in query_lower,
        'give me an overview' in query_lower,
        'brief overview' in query_lower
    ]) and not any([
        'how to implement' in query_lower,
        'how to achieve' in query_lower,
        'how should we' in query_lower,
        'help us implement' in query_lower,
        'guide us' in query_lower,
        'steps to' in query_lower,
        'how do i' in query_lower,
        'how can i' in query_lower
    ])
    
    # Create a comprehensive synthesis prompt
    expert_analyses_text = ""
    for i, output in enumerate(outputs, 1):
        expert_analyses_text += f"\n--- Expert Analysis {i} ---\n{output}\n"
    
    if is_concise:
        prompt = f"""
Synthesize these expert analyses into a CONCISE, actionable response.

Previous conversation context:
{conversation_context}

Original Query: {query}
Expert Analyses:{expert_analyses_text[:1000]}

**FOLLOW-UP HANDLING:** If this is a follow-up question, reference the previous conversation context above.

**CRITICAL OUTPUT FORMAT - HTML REQUIRED:**
Format your response in valid HTML. Use:
- <h2> for headings
- <ul><li> for bullet points
- <p> for paragraphs
- Do NOT use markdown (no ##, *, etc.) - use HTML only

Provide response as:
<h2>Key Steps</h2>
<ul>
<li>Action 1 (cite evidence if from documents)</li>
<li>Action 2</li>
<li>Action 3</li>
</ul>

<h2>Critical Requirements</h2>
<ul>
<li>Requirement 1 (cite evidence if from documents)</li>
<li>Requirement 2</li>
</ul>

**MANDATORY HIGHLIGHTING RULES:**
- ALL control IDs, article numbers, and framework references MUST be highlighted in green: <span style=\"color:#008000\">Article 17 GDPR</span>, <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>
- Evidence quotes from documents MUST be highlighted: <span style=\"color:#008000\">\"exact quote\"</span>
- Example: Data subjects have the right to erasure (<span style=\"color:#008000\">Article 17 GDPR</span>).

Keep total response under 200 words. Focus only on actionable steps. Output in HTML format.
"""
    elif is_simple_info_query:
        # For simple informational queries, provide a clean, direct response
        prompt = f"""
You are a compliance expert explaining a compliance concept or framework to a professional audience.

Previous conversation context:
{conversation_context}

Original Query: {query}

Multiple expert perspectives on this topic:{expert_analyses_text}

**CRITICAL INSTRUCTIONS:**
**FOLLOW-UP HANDLING:** If this is a follow-up question, review the previous conversation context above and reference it when relevant.
Synthesize the information above into ONE unified, cohesive explanation. DO NOT reference "Expert 1", "Expert 2", or compare different analyses. Present the information as if you are the single authoritative source.

Structure your response with:

1. **Overview**: Start with a clear definition and purpose of the framework/concept
2. **Key Components**: Explain the main components, rules, or requirements
3. **Who It Applies To**: Covered entities, organizations, or industries affected
4. **Main Requirements**: Core compliance obligations and standards
5. **Key Takeaways**: Important points to remember

**CRITICAL OUTPUT FORMAT - HTML REQUIRED:**
Format your entire response in valid HTML. Use:
- <h2> for main headings, <h3> for subheadings
- <ul><li> for bullet points
- <p> for paragraphs
- <strong> for bold, <em> for italic
- Do NOT use markdown (no ##, *, etc.) - use HTML only

Guidelines:
- Merge all expert insights into a single, unified voice
- DO NOT mention "experts", "analyses", "perspectives", or compare viewpoints
- Use clear, professional language as if this is YOUR direct knowledge
- Eliminate all redundancy and present information once
- Keep it informative but accessible
- Focus on understanding the concept, not implementation details
- Use a logical flow from general to specific
- **MANDATORY HIGHLIGHTING RULES:**
  * ALL control IDs, article numbers, and framework references MUST be highlighted in green: <span style=\"color:#008000\">Article 17 GDPR</span>, <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>, <span style=\"color:#008000\">HIPAA §164.312</span>
  * Evidence quotes from documents MUST be highlighted: <span style=\"color:#008000\">\"exact quote\"</span>
  * Example: Organizations must implement access controls (<span style=\"color:#008000\">ISO 27001 A.9.2.1</span>).

Provide a comprehensive, unified explanation in HTML format:
"""
    else:
        # For implementation/how-to queries, use the Gatekeeper format
        prompt = f"""
You are a Senior Compliance Architect acting as a Quality Gatekeeper and Technical Lead.

Previous conversation context:
{conversation_context}

Original Query: {query}
Context: {context}

Expert Analyses to Synthesize:{expert_analyses_text}

**INPUT ANALYSIS:**
User's Technology Stack: [DETECT FROM CONTEXT: e.g., AWS, Azure, GCP]

**CRITICAL SYNTHESIS INSTRUCTIONS (The "Gatekeeper" Rules):**
1. **Sanity Check the Tech Stack:**
   - Review the "Expert Analyses".
   - IF the user uses AWS and an expert recommends Azure-specific tooling, REMOVE/replace it with the equivalent AWS control (e.g., replace "Azure Sentinel" with "AWS Security Hub") or keep it vendor-neutral.
   - Do not allow hallucinated tools or mismatched vendors into the final response.
2. **Conflict Resolution:**
   - If Privacy guidance says "Delete Data" and Audit guidance says "Retain Logs", explain how to balance both (e.g., retain logs but anonymize personal data).
3. **Citation & Highlight Preservation (CRITICAL):**
   - Retain the specific Article numbers / Control IDs supplied by experts (e.g., GDPR Article 17, ISO 27001 A.9.2.1).
   - ALWAYS preserve green highlighting for ALL document-derived evidence: <span style=\"color:#008000\">\"exact quote from framework\"</span>
   - **MANDATORY:** ALL control IDs, article numbers, and framework references MUST be highlighted in green containers, whether from documents or general knowledge:
     * <span style=\"color:#008000\">Article 17 GDPR</span>
     * <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>
     * <span style=\"color:#008000\">HIPAA §164.312(a)(1)</span>
     * <span style=\"color:#008000\">PCI DSS Requirement 1.1</span>
   - Example format: Data subjects have the right to erasure ("right to be forgotten") (<span style=\"color:#008000\">Article 17 GDPR</span>).
   - When experts cite evidence from documents, ensure it appears in green in the final response.
   - Any mention of framework requirements, control IDs, or regulation articles MUST be highlighted in green, regardless of source.

**FINAL RESPONSE STRUCTURE:**
1. **Executive Summary**: A clear, direct answer to the user's strategy.
2. **Key Compliance & Control Mapping**:
   **CRITICAL - THIS SECTION MUST USE HTML TABLE TAGS (NOT MARKDOWN):**
   - For the "Key Compliance & Control Mapping" section ONLY, you MUST use HTML table tags
   - DO NOT use markdown table syntax (no | separators, no |---|---|)
   - **DO NOT wrap the HTML table in markdown code blocks** (no ```html ... ```)
   - Output the HTML table tags directly in your response
   - The table will have proper borders and styling applied automatically
   - Use this EXACT HTML structure:
   
   <table>
     <thead>
       <tr>
         <th>Regulation/Control ID</th>
         <th>Requirement</th>
         <th>Specific Tool/Action in User's Tech Stack</th>
       </tr>
     </thead>
     <tbody>
       <tr>
         <td><span style="color:#008000">Article 17 GDPR</span></td>
         <td>Right to erasure ("right to be forgotten")</td>
         <td>Implement automated deletion using AWS S3 Lifecycle Policies</td>
       </tr>
       <tr>
         <td><span style="color:#008000">ISO 27001 A.9.2.1</span></td>
         <td>User access management</td>
         <td>Implement strict IAM roles and policies following least privilege</td>
       </tr>
     </tbody>
   </table>
   
   **Requirements for this table:**
   - Use <table>, <thead>, <tbody>, <tr>, <th>, <td> tags
   - DO NOT use markdown table syntax (no | separators)
   - DO NOT wrap in markdown code blocks (no ```html or ```)
   - Output raw HTML table tags directly
   - Each row must be a <tr> with <td> cells inside
   - Control IDs and Article numbers MUST be in green: <span style="color:#008000">Article 17 GDPR</span>
   - All three columns are required: Regulation/Control ID, Requirement, Specific Tool/Action
   - Preserve green-highlighted evidence for any direct quotes
   - The table will automatically have borders and proper styling - you don't need to add inline styles

3. **Detailed Implementation Plan (SCOPE-ADAPTIVE)**:
   - **CRITICAL FORMATTING RULE:** Adopt the structure that best fits the User's Query Scope.
   - **Scenario A (Broad Process):** If the user asks for a full workflow (e.g., "How do I handle data?", "GDPR Guide"), you MUST use the Data Lifecycle stages (Collection → Processing → Storage → Deletion).
   - **Scenario B (Specific Control):** If the user asks for a specific topic (e.g., "How to encrypt", "Audit Logs"), structure the response by Configuration Steps (e.g., Planning → Configuration → Verification).
   - **Constraint:** Do NOT force a lifecycle structure on a narrow technical question.
4. **Risk & Governance**: Critical risks and specific mitigation steps.
5. **Evidence for Auditors** (if applicable): When listing evidence items, artifacts, or documentation requirements:
   - **CRITICAL FORMATTING:** Format as a numbered list with each item on a separate line.
   - Use HTML ordered list tags: <ol><li>Item 1</li><li>Item 2</li></ol> OR format with line breaks between numbered items.
   - Each evidence item must be on its own line, NOT in paragraph form.
   - Example of CORRECT format:
     <ol>
     <li>Backup Policy: A documented backup policy that defines the scope, frequency, retention, and testing procedures.</li>
     <li>Backup Configuration: Configuration details of the backup systems, including schedules and storage locations.</li>
     <li>Backup Logs: Logs showing that backups are being performed according to the defined schedule.</li>
     </ol>
   - Example of INCORRECT format (DO NOT USE):
     "1. Backup Policy: ... 2. Backup Configuration: ... 3. Backup Logs: ..." (all in one paragraph)

ADDITIONAL RULES:
- DO NOT say 'not in documents', 'not available', 'there is no direct mention', 'the document focuses on', or explain what's missing. Simply provide the information.
- DO NOT explain your approach, methodology, or what you're doing. Just provide the answer directly.
- DO NOT include phrases like "I will provide", "I will proceed", "focusing on a two-part approach", or similar meta-commentary.
- Reference previous conversation details when this is a follow-up.
- Remove redundancy; present one aligned strategy in a single authoritative voice.
- **MANDATORY:** ALL control IDs, article numbers, regulation sections, and framework references MUST be highlighted in green containers: <span style=\"color:#008000\">Article 17 GDPR</span>, <span style=\"color:#008000\">ISO 27001 A.9.2.1</span>, <span style=\"color:#008000\">HIPAA §164.312</span>
- **MANDATORY:** All evidence quotes from framework documents MUST be highlighted in green: <span style=\"color:#008000\">\"exact quote\"</span>
- Example: Organizations must implement access controls (<span style=\"color:#008000\">ISO 27001 A.9.2.1</span>).
- Example: Data subjects have the right to erasure ("right to be forgotten") (<span style=\"color:#008000\">Article 17 GDPR</span>).
- Clearly distinguish document-based evidence (with green highlight + citation) from supplemental guidance.
- When mentioning Control IDs, Article numbers, or regulation sections, ALWAYS wrap them in green highlighting, regardless of whether they came from documents or general knowledge.

Generate a unified, technically accurate, vendor-consistent response without meta-commentary:
"""
    
    return rate_limited_generate_content_optimized(prompt, max_tokens=max_tokens)

# Add ConversationHistory class to maintain context across queries
class ConversationHistory:
    def __init__(self, max_turns=3, timeout_seconds=600):
        self.history = []
        self.max_turns = max_turns
        self.timeout_seconds = timeout_seconds
        self.last_update_time = time.time()
        self.context_embedding = None
        self.last_compliance_status = True  # Track if last query was compliance-related

    def add_exchange(self, user_query, bot_response, is_compliance=True):
        current_time = time.time()
        if current_time - self.last_update_time > self.timeout_seconds:
            self.reset()
        
        # Only add to history if it's a compliance-related query
        if is_compliance:
            # Limit response size to prevent context growth
            truncated_response = bot_response
            if len(truncated_response) > 1000:
                truncated_response = truncated_response[:950] + "... [truncated for brevity]"
            
            self.history.append((user_query, truncated_response))
            if len(self.history) > self.max_turns:
                self.history.pop(0)
            
            self.last_update_time = current_time
            self.context_embedding = None  # Reset - will be lazily generated when needed
        
        self.last_compliance_status = is_compliance

    def get_context(self, compact=False):
        # If last query was non-compliance, don't provide context
        if not self.last_compliance_status:
            return ""
            
        if not self.history:
            return ""
        
        if compact:
            context = ""
            for i, (user, bot) in enumerate(self.history[-2:]):
                context += f"User: {user}\nResponse: [Summary of previous response]\n"
            return context
        
        context = "Previous conversation:\n"
        for user, bot in self.history:
            context += f"User: {user}\nBot: {bot}\n"
        return context

    def reset(self):
        """Reset conversation history"""
        self.history = []

    

# Track recent non-compliance responses to ensure variety
recent_non_compliance_responses = {}

def generate_non_compliance_response(query: str, user_id: str = None) -> str:
    """Generate a dynamic, contextual response for non-compliance queries"""
    try:
        # Analyze the query to understand what the user is asking about
        query_lower = query.lower()
        
        # Determine the category of non-compliance query
        response_category = "general"
        query_subject = ""
        
        # Categorize the query
        if any(word in query_lower for word in ["what is", "what are", "explain", "define"]):
            if any(word in query_lower for word in ["cooking", "recipe", "food", "restaurant"]):
                response_category = "cooking_food"
                query_subject = "cooking and food topics"
            elif any(word in query_lower for word in ["movie", "film", "tv", "show", "entertainment"]):
                response_category = "entertainment"
                query_subject = "entertainment topics"
            elif any(word in query_lower for word in ["sports", "football", "cricket", "basketball", "game"]):
                response_category = "sports"
                query_subject = "sports topics"
            elif any(word in query_lower for word in ["health", "medical", "doctor", "medicine"]):
                response_category = "health"
                query_subject = "personal health topics"
            elif any(word in query_lower for word in ["weather", "climate", "temperature"]):
                response_category = "weather"
                query_subject = "weather information"
            else:
                response_category = "general_knowledge"
                query_subject = "general knowledge topics"
        elif any(word in query_lower for word in ["how to", "tutorial", "guide"]):
            response_category = "tutorial"
            query_subject = "tutorials and guides"
        elif any(word in query_lower for word in ["joke", "funny", "humor"]):
            response_category = "humor"
            query_subject = "humor and jokes"
        
        # Check if user has recent responses to avoid repetition
        avoid_phrases = []
        if user_id and user_id in recent_non_compliance_responses:
            recent_responses = recent_non_compliance_responses[user_id]
            # Extract key phrases from recent responses to avoid
            for recent_response in recent_responses[-3:]:  # Check last 3 responses
                if "curious about" in recent_response:
                    avoid_phrases.append("curious about")
                if "interesting question" in recent_response:
                    avoid_phrases.append("interesting question")
                if "outside my area" in recent_response:
                    avoid_phrases.append("outside my area")
        
        # Generate a contextual response using AI
        prompt = f"""
Generate a polite, helpful response for a user who asked a non-compliance question in a compliance chatbot.

User's question: "{query}"
Question category: {response_category}
Question subject: {query_subject}
Avoid using these phrases: {', '.join(avoid_phrases) if avoid_phrases else 'none'}

The response should:
1. Be polite and understanding
2. Briefly acknowledge what they're asking about
3. Redirect them to compliance topics
4. Suggest 2-3 specific compliance topics that might interest them
5. Be conversational and helpful, not robotic
6. Be 2-3 sentences maximum
7. Use different phrasing than previous responses

Examples of compliance topics to suggest:
- Data protection and GDPR compliance
- Security frameworks (ISO 27001, SOC 2)
- Identity and access management
- Cloud security compliance
- Audit requirements and controls
- Privacy policy development
- Risk assessment procedures
- Vendor compliance management

Make it sound natural and varied - don't use the same phrases every time.
Avoid being overly formal or robotic.
"""

        response = rate_limited_generate_content_optimized(prompt, max_tokens=3200)  # Increased max_tokens for more variety
        
        # Add compliance topic suggestions based on query context
        topic_suggestions = get_contextual_compliance_suggestions(query_lower, response_category)
        
        if topic_suggestions and response:
            response += f" {topic_suggestions}"
        
        # Fallback responses if AI generation fails
        if not response or len(response.strip()) < 10:
            fallback_responses = [
                f"I understand you're curious about {query_subject if query_subject else 'that topic'}, but I'm specialized in compliance and regulatory matters. I'd be happy to help you with questions about data protection, security frameworks, or audit requirements instead!",
                
                f"That's an interesting question about {query_subject if query_subject else 'that subject'}! However, my expertise is in compliance and governance. Feel free to ask me about GDPR compliance, identity management, or security controls.",
                
                f"I can see you're asking about {query_subject if query_subject else 'that area'}, but I'm designed to assist with compliance topics. I'd love to help you understand privacy regulations, risk assessments, or compliance frameworks instead!",
                
                f"While that's outside my area of expertise, I'm here to help with all things compliance! Whether you need information about ISO 27001, data protection policies, or vendor compliance, I'm ready to assist.",
                
                f"I focus specifically on compliance, security, and regulatory guidance. Though I can't help with {query_subject if query_subject else 'that topic'}, I'd be excited to discuss cloud security, audit procedures, or privacy impact assessments with you!",
                
                f"My specialty is in regulatory and compliance matters rather than {query_subject if query_subject else 'that area'}. How about we explore some fascinating compliance topics like zero-trust security, privacy by design, or regulatory change management?",
                
                f"That question falls outside my compliance expertise, but I'd love to help you navigate the world of data governance, security compliance, or regulatory frameworks instead!",
                
                f"I'm built to tackle compliance challenges! While I can't assist with {query_subject if query_subject else 'that topic'}, I'm excited to discuss compliance automation, risk management strategies, or security audit procedures with you.",
                
                f"Ah, {query_subject if query_subject else 'that topic'} isn't in my wheelhouse, but compliance definitely is! How about we dive into some exciting areas like breach response planning, vendor risk assessments, or privacy impact analyses?",
                
                f"I wish I could help with {query_subject if query_subject else 'that'}, but my superpower is in compliance and security! Let's talk about something in my domain - perhaps regulatory reporting, access controls, or compliance monitoring?",
                
                f"That's not really my thing, but compliance frameworks absolutely are! I'd love to chat about topics like SOC 2 implementations, GDPR compliance strategies, or security control effectiveness instead.",
                
                f"I'm afraid {query_subject if query_subject else 'that topic'} isn't where I shine, but ask me anything about compliance and I'm your bot! How about we explore business continuity planning, third-party risk management, or regulatory change tracking?",
                
                f"Sorry, but {query_subject if query_subject else 'that area'} is outside my zone of expertise. I'm much better with compliance topics like incident response procedures, control testing methodologies, or privacy program development!",
                
                f"I'd love to help, but {query_subject if query_subject else 'that'} isn't my forte. My strength lies in compliance matters - shall we discuss compliance training programs, audit readiness, or regulatory gap analyses instead?"
            ]
            
            # Filter out responses with avoided phrases
            if avoid_phrases:
                filtered_responses = [r for r in fallback_responses if not any(phrase in r.lower() for phrase in avoid_phrases)]
                fallback_responses = filtered_responses if filtered_responses else fallback_responses
            
            response = random.choice(fallback_responses)
        
        # Track this response for variety in future interactions
        if user_id:
            if user_id not in recent_non_compliance_responses:
                recent_non_compliance_responses[user_id] = []
            recent_non_compliance_responses[user_id].append(response)
            # Keep only last 5 responses to avoid memory issues
            if len(recent_non_compliance_responses[user_id]) > 5:
                recent_non_compliance_responses[user_id].pop(0)
        
        return response.strip()
        
    except Exception as e:
        logger.error(f"Error generating non-compliance response: {e}")
        # Final fallback
        return "I'm a compliance assistant focused on regulatory and security topics. I'd be happy to help you with compliance frameworks, privacy regulations, or security requirements instead!"

def generate_simple_non_compliance_response(query: str) -> str:
    """Generate a simple, concise non-compliance response (1-2 sentences max)"""
    try:
        simple_responses = [
            "I'm a compliance assistant focused on privacy, security, and regulatory topics. Feel free to ask about GDPR, ISO 27001, privacy policies, or other compliance matters!",
            "That's outside my area of expertise. I specialize in compliance frameworks, data protection, and security regulations. How can I help with those?",
            "I focus on compliance and regulatory guidance. Ask me about privacy policies, security frameworks, or compliance requirements!",
            "I'm here to help with compliance topics like GDPR, HIPAA, ISO 27001, and security best practices. What compliance question can I answer for you?",
            "That's not a compliance topic I can help with. I'd be happy to discuss data protection, security frameworks, or regulatory requirements instead!",
            "I specialize in compliance, privacy, and security matters. Feel free to ask about frameworks, regulations, or policy development!"
        ]
        return random.choice(simple_responses)
    except Exception as e:
        logger.error(f"Error generating simple non-compliance response: {e}")
        return "I'm a compliance assistant focused on regulatory and security topics. How can I help with compliance matters?"

def get_contextual_compliance_suggestions(query_lower: str, category: str) -> str:
    """Generate contextual compliance topic suggestions based on the query"""
    try:
        suggestions_map = {
            "cooking_food": "You might be interested in food industry compliance, restaurant data protection policies, or PCI DSS for payment processing.",
            "entertainment": "Perhaps you'd like to know about content privacy policies, digital media compliance, or COPPA regulations for entertainment platforms?",
            "sports": "You could explore sports data analytics compliance, fan data protection, or privacy policies for sports platforms.",
            "health": "I can help with HIPAA compliance, healthcare data protection, or medical device security standards instead.",
            "weather": "Maybe you're interested in IoT device compliance, environmental data protection, or smart city privacy frameworks?",
            "tutorial": "I can guide you through compliance implementation tutorials, security assessment procedures, or privacy policy creation guides.",
            "humor": "How about some 'seriously fun' compliance topics like gamification in security training or user-friendly privacy notices?",
            "general_knowledge": "I can share knowledge about regulatory frameworks, compliance best practices, or security standards."
        }
        
        base_suggestions = [
            "Would you like to explore identity and access management or cloud security instead?",
            "How about learning about data protection laws or security compliance frameworks?",
            "I could help you understand audit requirements or privacy regulations.",
            "Perhaps you're interested in risk assessment procedures or compliance monitoring?",
            "Would GDPR compliance, ISO 27001, or SOC 2 frameworks be helpful to discuss?"
        ]
        
        if category in suggestions_map:
            return suggestions_map[category]
        else:
            return random.choice(base_suggestions)
            
    except Exception as e:
        logger.error(f"Error generating compliance suggestions: {e}")
        return "Feel free to ask about any compliance, security, or regulatory topics!"

def search_documents(query: str, segments: List[str], index: Any, top_k: int = 3) -> List[str]:
    """
    Search documents using FAISS similarity search.
    
    Args:
        query (str): Search query
        segments (List[str]): List of document segments
        index (Any): FAISS index
        top_k (int): Number of top results to return
        
    Returns:
        List[str]: List of relevant segments
    """
    try:
        # Get query embedding
        query_embedding = get_embedding(query)
        if query_embedding is None:
            return []
        
        # Search using FAISS index
        query_embedding = np.expand_dims(query_embedding, axis=0)
        distances, idxs = index.search(query_embedding, top_k)
        
        # Return relevant segments
        relevant_segments = []
        for idx in idxs[0]:
            if idx < len(segments):
                relevant_segments.append(segments[idx])
        
        return relevant_segments
        
    except Exception as e:
        logger.error(f"Error searching documents: {e}")
        return []

# Precomputed expert selection for performance
EXPERT_KEYWORD_SCORES = {
    'security': {
        'azure': 5, 'azure ad': 5, 'identity': 4, 'authentication': 4, 'access control': 4,
        'security': 3, 'firewall': 3, 'encryption': 3, 'cybersecurity': 3
    },
    'privacy': {
        'gdpr': 5, 'ccpa': 5, 'privacy': 4, 'data protection': 4, 'personal data': 4,
        'consent': 3, 'data subject': 3, 'pii': 3
    },
    'audit': {
        'compliance': 4, 'audit': 4, 'iso27001': 5, 'soc2': 5, 'nist': 4,
        'framework': 3, 'control': 3, 'assessment': 3
    },
    'financial': {
        'pci dss': 5, 'sox': 5, 'financial': 4, 'payment': 4, 'banking': 3
    },
    'healthcare': {
        'hipaa': 5, 'healthcare': 4, 'phi': 4, 'protected health information': 4,
        'health insurance portability': 3, 'medical': 3, 'patient data': 3,
        'health data': 3, 'ehr': 3, 'electronic health record': 3, 'ephi': 3,
        'covered entity': 3, 'business associate': 3
    }
}

@timing_decorator
def select_relevant_experts_optimized(query: str) -> List[str]:
    """Optimized expert selection using precomputed scores."""
    query_lower = query.lower()
    expert_scores = {}
    
    # Fast scoring using precomputed weights
    for expert, keywords in EXPERT_KEYWORD_SCORES.items():
        score = 0
        for keyword, weight in keywords.items():
            if keyword in query_lower:
                score += weight
        if score > 0:
            expert_scores[expert] = score
    
    # Return top 2 experts for speed (reduced from 3)
    if expert_scores:
        sorted_experts = sorted(expert_scores.items(), key=lambda x: x[1], reverse=True)
        selected = [expert for expert, score in sorted_experts[:2]]
        logger.info(f"Fast expert selection: {', '.join(selected)}")
        return selected
    
    # Quick fallback
    return ['audit', 'security']

    

# Modify analyze_privacy_policy to accept a specific framework for analysis
def analyze_privacy_policy(file_path: str, segments: List[str], index: Any, framework: str) -> str:
    """
    Analyze a privacy policy document against a specific compliance framework.
    """
    try:
        # Extract text from the document
        if file_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        else:
            return "Unsupported file format. Please upload a PDF or DOCX file."

        # Get framework requirements
        requirements = get_framework_requirements(framework)
        
        # Analyze each requirement
        analysis_results = []
        for req in requirements:
            # Search for relevant content in the document
            query = f"Find content related to {req['title']} and {req['description']}"
            relevant_segments = search_documents(query, segments, index, top_k=3)
            
            # Check if requirement is met
            is_met = any(req['title'].lower() in segment.lower() or 
                        req['description'].lower() in segment.lower() 
                        for segment in relevant_segments)
            
            analysis_results.append({
                'requirement': req['title'],
                'description': req['description'],
                'is_met': is_met,
                'relevant_content': relevant_segments if is_met else []
            })
        
        # Generate analysis report
        report = f"Privacy Policy Analysis against {framework} Requirements\n\n"
        
        # Summary
        total_reqs = len(analysis_results)
        met_reqs = sum(1 for r in analysis_results if r['is_met'])
        report += f"Summary:\n"
        report += f"- Total Requirements: {total_reqs}\n"
        report += f"- Requirements Met: {met_reqs}\n"
        report += f"- Requirements Missing: {total_reqs - met_reqs}\n"
        report += f"- Compliance Score: {(met_reqs/total_reqs)*100:.1f}%\n\n"
        
        # Detailed Analysis
        report += "Detailed Analysis:\n"
        for result in analysis_results:
            report += f"\nRequirement: {result['requirement']}\n"
            report += f"Description: {result['description']}\n"
            report += f"Status: {'✓ Met' if result['is_met'] else '✗ Missing'}\n"
            if result['is_met']:
                report += "Relevant Content Found:\n"
                for content in result['relevant_content']:
                    report += f"- {content}\n"
            else:
                report += "Recommendation: Add content addressing this requirement\n"
        
        return report
        
    except Exception as e:
        logger.error(f"Error analyzing privacy policy: {str(e)}")
        return f"Error analyzing privacy policy: {str(e)}"

# Add function to generate a privacy policy document
def generate_privacy_policy(framework: str, format: str = "txt") -> str:
    """
    Generate a privacy policy document based on a specific compliance framework.
    """
    try:
        # Get framework requirements
        requirements = get_framework_requirements(framework)
        
        # Generate policy sections
        sections = []
        
        # Introduction
        sections.append(f"Privacy Policy\n")
        sections.append(f"Last Updated: {datetime.now().strftime('%Y-%m-%d')}\n")
        sections.append(f"This Privacy Policy describes how we collect, use, and protect your personal information in compliance with {framework} requirements.\n")
        
        # Generate content for each requirement using AI
        for req in requirements:
            section_title = req['title']
            prompt = (
                f"Generate a detailed section for a {framework} compliant privacy policy about {req['title']}.\n"
                f"Requirement: {req['description']}\n"
                "The section should:\n"
                "1. Be specific and actionable\n"
                "2. Include all necessary legal requirements\n"
                "3. Use clear, professional language\n"
                "4. Be comprehensive but concise\n"
                "5. Include practical implementation details\n\n"
                "Generate the section content:"
            )
            
            section_content = rate_limited_generate_content(prompt)
            sections.append(f"\n{section_title}\n")
            sections.append(f"{section_content}\n")
        
        # Add standard sections
        standard_sections = {
            "Contact Information": "For any questions or concerns regarding this Privacy Policy or our data practices, please contact us at:\n\n[Company Name]\n[Address]\n[Email]\n[Phone]",
            "Changes to This Policy": f"We may update this Privacy Policy from time to time to reflect changes in our practices or for other operational, legal, or regulatory reasons. The updated version will be indicated by an updated 'Last Updated' date.",
            "Compliance and Certification": f"We are committed to maintaining compliance with {framework} requirements and regularly review our practices to ensure they meet the highest standards of data protection and privacy."
        }
        
        for title, content in standard_sections.items():
            sections.append(f"\n{title}\n")
            sections.append(f"{content}\n")
        
        # Combine all sections
        policy_text = "\n".join(sections)
        
        # Convert to requested format
        if format.lower() == "pdf":
            # Convert to PDF
            doc = Document()
            for section in sections:
                doc.add_paragraph(section)
            pdf_path = f"generated_policy_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            doc.save(pdf_path)
            return f"Policy generated and saved as {pdf_path}"
            
        elif format.lower() == "docx":
            # Convert to DOCX
            doc = Document()
            for section in sections:
                doc.add_paragraph(section)
            docx_path = f"generated_policy_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(docx_path)
            return f"Policy generated and saved as {docx_path}"
            
        else:
            return policy_text
            
    except Exception as e:
        logger.error(f"Error generating privacy policy: {str(e)}")
        return f"Error generating privacy policy: {str(e)}"

def get_framework_requirements(framework: str) -> List[Dict[str, str]]:
    """
    Get the requirements for a specific compliance framework.
    """
    requirements = {
        "GDPR": [
            {
                "title": "Data Collection and Processing",
                "description": "Clear information about what personal data is collected, how it is processed, and the purposes of processing"
            },
            {
                "title": "Legal Basis for Processing",
                "description": "Explanation of the legal grounds for processing personal data, including consent, contract, legal obligation, vital interests, public task, and legitimate interests"
            },
            {
                "title": "Data Subject Rights",
                "description": "Detailed explanation of data subject rights under GDPR, including right to access, rectification, erasure, restriction, portability, and objection"
            },
            {
                "title": "Data Retention",
                "description": "Clear information about data retention periods and criteria for determining retention periods"
            },
            {
                "title": "Data Security",
                "description": "Comprehensive description of technical and organizational security measures implemented to protect personal data"
            },
            {
                "title": "International Data Transfers",
                "description": "Information about international data transfers and appropriate safeguards implemented"
            },
            {
                "title": "Data Breach Notification",
                "description": "Procedures for handling and reporting data breaches in compliance with GDPR requirements"
            },
            {
                "title": "Data Protection Officer",
                "description": "Information about the Data Protection Officer (DPO) and their contact details"
            }
        ],
        "CCPA": [
            {
                "title": "Personal Information Collection",
                "description": "Categories of personal information collected and the business or commercial purposes for collection"
            },
            {
                "title": "Consumer Rights",
                "description": "Detailed explanation of consumer rights under CCPA, including right to know, delete, and opt-out"
            },
            {
                "title": "Data Sales",
                "description": "Information about sale of personal information and opt-out rights"
            },
            {
                "title": "Financial Incentives",
                "description": "Details about any financial incentives offered for personal information"
            }
        ],
        "HIPAA": [
            {
                "title": "Protected Health Information",
                "description": "How PHI is collected, used, and disclosed in compliance with HIPAA requirements"
            },
            {
                "title": "Privacy Practices",
                "description": "Detailed description of privacy practices and patient rights under HIPAA"
            },
            {
                "title": "Security Measures",
                "description": "Comprehensive description of security measures implemented to protect PHI"
            }
        ]
    }
    
    return requirements.get(framework, requirements["GDPR"])

def generate_section_content(requirement: str, framework: str) -> str:
    """
    Generate content for a specific section based on the requirement and framework.
    """
    # This is a simplified version. In a real implementation, you would use
    # a more sophisticated approach, possibly using a language model to generate
    # appropriate content based on the requirement and framework.
    
    content_templates = {
        "GDPR": {
            "Data Collection and Processing": "We collect and process your personal data in accordance with GDPR requirements. This includes [specific data types] which we use for [specific purposes].",
            "Legal Basis": "We process your personal data based on the following legal grounds: [list of legal bases].",
            "Data Subject Rights": "Under GDPR, you have the right to access, rectify, erase, and port your personal data. You can exercise these rights by [instructions].",
            "Data Retention": "We retain your personal data for [specific period] or as required by law.",
            "Data Security": "We implement appropriate technical and organizational measures to protect your personal data."
        },
        "CCPA": {
            "Personal Information Collection": "We collect the following categories of personal information: [list of categories].",
            "Business Purpose": "We collect your personal information for the following business purposes: [list of purposes].",
            "Consumer Rights": "Under CCPA, you have the right to know, delete, and opt-out of the sale of your personal information.",
            "Data Sales": "We [do/do not] sell your personal information to third parties."
        },
        "HIPAA": {
            "Protected Health Information": "We collect and use your protected health information (PHI) for treatment, payment, and healthcare operations.",
            "Privacy Practices": "We maintain the privacy of your PHI and provide you with notice of our privacy practices.",
            "Security Measures": "We implement appropriate safeguards to protect your PHI from unauthorized access, use, or disclosure."
        }
    }
    
    # Find the most relevant template
    for key, template in content_templates.get(framework, content_templates["GDPR"]).items():
        if key.lower() in requirement.lower():
            return template
    
    return f"We handle {requirement.lower()} in accordance with {framework} requirements."

# Cache embeddings for performance
@lru_cache(maxsize=1000)
def get_cached_embedding(text_hash: str, text: str) -> np.ndarray:
    """Cached version of get_embedding"""
    return get_embedding(text)

# Hash function for text to use with lru_cache
def hash_text(text: str) -> str:
    """Create a hash of text for caching purposes"""
    return hashlib.md5(text.encode()).hexdigest()

# Add cached version of expertise functions
def cached_expert_response(expert_type: str, query: str, context: str, conversation_context: str = "") -> str:
    """Get expert response from cache if available, otherwise generate and cache it"""
    cache_key = f"{expert_type}:{hash_text(query)}:{hash_text(context)}:{hash_text(conversation_context)}"
    
    if cache_key in QUERY_CACHE:
        logger.info(f"Cache hit for {expert_type} expert")
        return QUERY_CACHE[cache_key]
    
    # Call appropriate expert function based on type
    if expert_type == "general":
        from compliance_rag_intelligent import expert_general_compliance
        response = expert_general_compliance(query, conversation_context, "")
    elif expert_type == "security":
        response = expert_security_controls(query, context, conversation_context)
    elif expert_type == "privacy":
        response = expert_privacy_regulations(query, context, conversation_context)
    elif expert_type == "audit":
        response = expert_audit_compliance(query, context, conversation_context)
    elif expert_type == "financial":
        response = expert_financial_compliance(query, context, conversation_context)
    elif expert_type == "healthcare":
        response = expert_healthcare_compliance(query, context, conversation_context)
    else:
        return ""
    
    # Cache the response
    QUERY_CACHE[cache_key] = response
    
    # Periodically save the cache (every 10 new entries)
    if len(QUERY_CACHE) % 10 == 0:
        save_query_cache()
        
    return response

    

@timing_decorator
def is_compliance_related_optimized(query: str, conversation_context: str = "") -> Tuple[bool, str]:
    """
    Optimized compliance query classification with fast-path processing.
    """
    
    # Fast Path 1: Check cache first
    cache_key = f"classification:{hash_text(query)}"
    if cache_key in QUERY_CACHE:
        cached_result = QUERY_CACHE[cache_key]
        if isinstance(cached_result, dict) and 'is_compliance' in cached_result:
            return cached_result['is_compliance'], cached_result.get('reason', 'Cached result')
    
    # Fast Path 2: Simple keyword screening for obviously non-compliant topics
    sensitive_topics = [
        "sex", "sexual", "porn", "pornography", "nude", "nudity",
        "drug", "drugs", "weapon", "weapons", "violence", "violent",
        "suicide", "self-harm", "cooking", "recipe", "sports", "football", 
        "cricket", "entertainment", "movie", "film", "music", "celebrity", "gossip"
    ]
    
    query_lower = query.lower()
    for topic in sensitive_topics:
        if topic in query_lower:
            professional_context = [
                "compliance", "policy", "regulation", "standard", "framework",
                "business", "organization", "company", "workplace", "professional"
            ]
            if not any(context in query_lower for context in professional_context):
                result = (False, f"Query about non-compliance topic: {topic}")
                QUERY_CACHE[cache_key] = {'is_compliance': False, 'reason': result[1]}
                return result
    
    # Fast Path 3: Strong compliance indicators
    strong_compliance_keywords = [
        "gdpr", "ccpa", "hipaa", "sox", "iso27001", "soc2", "nist", "pci dss",
        "data protection", "privacy policy", "compliance", "audit", "security framework",
        "regulatory", "governance", "risk management", "azure ad", "identity management"
    ]
    
    for keyword in strong_compliance_keywords:
        if keyword in query_lower:
            result = (True, f"Strong compliance keyword detected: {keyword}")
            QUERY_CACHE[cache_key] = {'is_compliance': True, 'reason': result[1]}
            return result
    
    # Fast Path 4: Business/technical context
    business_technical_keywords = [
        "azure", "cloud", "security", "identity", "authentication", "access control",
        "database", "network", "infrastructure", "application", "system",
        "organization", "business", "company", "enterprise", "corporate"
    ]
    
    if any(keyword in query_lower for keyword in business_technical_keywords):
        result = (True, f"Business/technical context detected")
        QUERY_CACHE[cache_key] = {'is_compliance': True, 'reason': result[1]}
        return result
    
    # Only use expensive AI analysis for borderline cases
    try:
        ai_result = analyze_query_intent_with_ai(query, conversation_context)
        confidence = ai_result.get('confidence', 0.0)
        # Check if intent suggests compliance-related query
        intent = ai_result.get('intent', 'GENERAL_COMPLIANCE')
        is_compliance_intent = intent in {'GENERAL_COMPLIANCE', 'SPECIFIC_REQUIREMENT', 'DOCUMENT_ANALYSIS', 'ANALYZE_UPLOADED'}
        
        result = (is_compliance_intent and confidence > 0.3, f"AI analysis: {confidence:.2f} confidence, intent: {intent}")
        QUERY_CACHE[cache_key] = {'is_compliance': result[0], 'reason': result[1]}
        return result
    except Exception as e:
        logger.warning(f"AI classification failed, defaulting to compliant: {e}")
        result = (True, "Defaulted to compliant due to analysis failure")
        QUERY_CACHE[cache_key] = {'is_compliance': True, 'reason': result[1]}
        return result

@timing_decorator
def detect_ambiguous_query(query: str, conversation_context: str = "") -> Tuple[bool, str]:
    """
    Detect if a query is ambiguous and requires clarification using LLM analysis.
    
    Args:
        query (str): The user's query
        conversation_context (str): The conversation context to help determine if query is ambiguous
        
    Returns:
        Tuple[bool, str]: (is_ambiguous, clarification_message)
        - is_ambiguous: True if query is ambiguous, False otherwise
        - clarification_message: Suggested clarification message if ambiguous, empty string otherwise
    """
    # Normalize query
    query_normalized = query.strip().lower()
    
    # FIRST: Check if query mentions a specific compliance framework
    # If it does, it's NOT ambiguous even if it's in a clarification loop
    framework_patterns = [
        r'\bgdpr\b', r'\biso\s*27001\b', r'\biso\s*27002\b', r'\biso/iec\s*27001\b', r'\biso/iec\s*27002\b',
        r'\bsoc\s*2\b', r'\bsoc2\b', r'\bsoc\s*ii\b',
        r'\bhipaa\b',
        r'\bpci\s*dss\b', r'\bpci-dss\b', r'\bpci\b',
        r'\bnist\b',
        r'\bccpa\b', r'\bcpra\b',
        r'\biso\s*22301\b', r'\biso\s*31000\b'
    ]
    
    mentions_framework = any(re.search(pattern, query_normalized, re.IGNORECASE) for pattern in framework_patterns)
    
    if mentions_framework:
        # Query mentions a framework - NOT ambiguous, proceed to answer
        logger.info(f"Query mentions framework: '{query}' - NOT ambiguous")
        return False, ""
    
    # Check if previous bot response was asking for clarification
    # If so, and current query is still vague, treat as ambiguous
    is_followup_to_clarification = False
    if conversation_context:
        context_lower = conversation_context.lower()
        # Check if bot's last response was asking for clarification
        clarification_indicators = [
            "could you please specify",
            "could you please provide",
            "please specify",
            "please provide",
            "what 'that thing' is",
            "what compliance framework",
            "which compliance framework",
            "what specific requirement",
            "what specific control",
            "i'd be happy to help",
            "i'd be happy to provide",
            "clarification"
        ]
        # Check if context contains bot asking for clarification
        if any(indicator in context_lower for indicator in clarification_indicators):
            # Check if current query is still vague (short, uses pronouns, lacks specifics)
            vague_patterns = [
                r'^(that|this|it|the thing|that thing|this thing|it|them|those|these)$',
                r'^(what|which|how|why|when|where)\s+(is|are|was|were|do|does|did|will|would|can|could|should|must)\s*(that|this|it|the thing|that thing)?\s*\?*$',
                r'^(tell me|explain|describe|what about|how about)\s*(that|this|it|the thing|that thing)?\s*\?*$',
            ]
            for pattern in vague_patterns:
                if re.match(pattern, query_normalized):
                    is_followup_to_clarification = True
                    logger.info(f"Detected vague follow-up to clarification request: '{query}'")
                    break
    
    # If this is a vague follow-up to a clarification request, treat as ambiguous
    if is_followup_to_clarification:
        return True, "I still need more specific information to help you. Could you please provide details such as:\n- Which compliance framework are you asking about? (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, etc.)\n- What specific requirement, control, or regulation?\n- What aspect would you like to know more about? (implementation, requirements, certification, etc.)"
    
    # Quick check: very short queries without context are likely ambiguous
    if len(query_normalized) < 10 and not conversation_context:
        return True, "I'd be happy to help! Could you please provide more details about what you'd like to know? For example:\n- What compliance framework are you interested in? (GDPR, ISO 27001, SOC 2, etc.)\n- What specific requirement or control do you need information about?\n- Are you looking for implementation guidance or regulatory requirements?"
    
    # Use LLM to intelligently detect ambiguous queries
    prompt = f"""Analyze this user query and determine if it is ambiguous or unclear. IMPORTANT: Consider the conversation context when making your decision.

Query: "{query}"
Conversation Context: "{conversation_context[:800] if conversation_context else 'None (this is a new conversation with no prior context)'}"

CRITICAL RULES:
1. **HIGHEST PRIORITY - Framework Mention**: If the query mentions a specific compliance framework (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, NIST, CCPA, etc.), it is NOT ambiguous - even if it's in a clarification loop. Examples: "yes gdpr in short", "all of gdpr main stuff", "tell me about ISO 27001" are NOT ambiguous.

2. **SPECIAL CASE - Clarification Loop Detection**: If the Conversation Context shows the bot previously asked for clarification (e.g., "Could you please specify...", "What compliance framework..."), and the current query is still vague (e.g., "that thing", "it", "this") AND doesn't mention a framework, then the query IS ambiguous - the user hasn't provided the requested clarification yet.

3. If Conversation Context contains relevant information (previous queries about frameworks, controls, requirements, etc.), then follow-up queries like "tell me more", "how do I do it", "what's required" are NOT ambiguous - they refer to the previous conversation.

4. A query is ONLY ambiguous if BOTH conditions are true:
   a) The query itself lacks specific details (e.g., "tell me more", "how do I do it")
   b) There is NO relevant conversation context to understand what the user is referring to
   c) The query does NOT mention a specific compliance framework

A query is AMBIGUOUS if:
1. It lacks specific details AND there is no conversation context (e.g., "tell me more" in a new conversation)
2. It uses vague pronouns or references AND there is no context to understand them (e.g., "what about that?" with no prior discussion)
3. It's too generic without mentioning frameworks, controls, regulations, or specific topics AND no context exists
4. It cannot be answered even with the conversation context

A query is NOT ambiguous if:
1. It mentions specific frameworks (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, etc.)
2. It asks about specific controls, articles, or requirements
3. It has enough context from the conversation history to understand what the user means
4. It's a clear follow-up question that references previous conversation (e.g., "tell me more" after discussing GDPR requirements)
5. The conversation context provides enough information to understand the query

Respond with ONLY a JSON object in this exact format:
{{
    "is_ambiguous": true/false,
    "reasoning": "brief explanation",
    "clarification_suggestion": "helpful question to ask the user (only if is_ambiguous is true, otherwise empty string)"
}}

Examples:
Query: "tell me more"
Conversation Context: "None"
Response: {{"is_ambiguous": true, "reasoning": "Query lacks specific topic or framework reference and there is no conversation context", "clarification_suggestion": "I'd be happy to provide more information! Could you please specify what you'd like to know more about? For example:\n- Which compliance framework? (GDPR, ISO 27001, SOC 2, etc.)\n- What specific requirement or control?\n- Are you looking for implementation steps or regulatory details?"}}

Query: "tell me more"
Conversation Context: "User asked: What are GDPR data breach notification requirements? Bot responded with information about GDPR Article 33 and 34 requirements."
Response: {{"is_ambiguous": false, "reasoning": "Query is a follow-up to previous GDPR discussion, context provides clear reference", "clarification_suggestion": ""}}

Query: "What are the GDPR data breach notification requirements?"
Conversation Context: "None"
Response: {{"is_ambiguous": false, "reasoning": "Query mentions specific framework (GDPR) and specific topic (data breach notification)", "clarification_suggestion": ""}}

Query: "how do i do it"
Conversation Context: "None"
Response: {{"is_ambiguous": true, "reasoning": "Query uses vague pronoun 'it' without context", "clarification_suggestion": "I'd be happy to help you! Could you please clarify what you'd like to do? For example:\n- Implement a specific compliance control?\n- Achieve certification for a framework?\n- Set up security measures?\n- Please provide more details about your goal."}}

Query: "how do i do it"
Conversation Context: "User asked: What are ISO 27001 access control requirements? Bot explained ISO 27001 Control A.9 requirements."
Response: {{"is_ambiguous": false, "reasoning": "Query refers to implementing ISO 27001 access controls from previous conversation", "clarification_suggestion": ""}}

Query: "what's required"
Conversation Context: "None"
Response: {{"is_ambiguous": true, "reasoning": "Query is too generic without specifying what is required and no context exists", "clarification_suggestion": "I can help you understand compliance requirements! Could you please specify:\n- Which compliance framework? (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, etc.)\n- What type of requirement? (access controls, encryption, data protection, audit, etc.)\n- For what purpose? (certification, implementation, assessment)"}}

Query: "what's required"
Conversation Context: "User asked about SOC 2 certification process. Bot explained the SOC 2 Trust Services Criteria."
Response: {{"is_ambiguous": false, "reasoning": "Query refers to SOC 2 requirements from previous conversation context", "clarification_suggestion": ""}}

Query: "that thing"
Conversation Context: "User: what is that thing tell me moree about it kindly. Bot: I'd be happy to provide more information! Could you please specify what 'that thing' is? For example, are you referring to a specific compliance framework, control, or regulation?"
Response: {{"is_ambiguous": true, "reasoning": "Bot previously asked for clarification about 'that thing', but user's response is still vague and doesn't provide the requested specifics or mention a framework", "clarification_suggestion": "I still need more specific information to help you. Could you please provide details such as:\n- Which compliance framework are you asking about? (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, etc.)\n- What specific requirement, control, or regulation?\n- What aspect would you like to know more about? (implementation, requirements, certification, etc.)"}}

Query: "yes gdpr in short"
Conversation Context: "User: i wanna know that. Bot: I'd be happy to help! Could you please specify what you'd like to know? For example, are you interested in a specific compliance framework (like GDPR, ISO 27001, or SOC 2)?"
Response: {{"is_ambiguous": false, "reasoning": "Query mentions specific framework (GDPR) and requests a short summary, which is a clear and answerable request", "clarification_suggestion": ""}}

Query: "all of gdpr main stuff"
Conversation Context: "User: dont u know what i am talking about. Bot: I still need more specific information to help you..."
Response: {{"is_ambiguous": false, "reasoning": "Query mentions specific framework (GDPR) and requests main/important information, which is a clear and answerable request despite being in a clarification loop", "clarification_suggestion": ""}}

Now analyze the query above and respond with ONLY the JSON object:"""

    try:
        response_text = rate_limited_generate_content_optimized(prompt, temperature=0.1, max_tokens=300)
        
        # Extract JSON from response - handle nested braces
        json_match = None
        # Try to find JSON object with proper brace matching
        brace_count = 0
        start_idx = -1
        for i, char in enumerate(response_text):
            if char == '{':
                if brace_count == 0:
                    start_idx = i
                brace_count += 1
            elif char == '}':
                brace_count -= 1
                if brace_count == 0 and start_idx != -1:
                    json_match = response_text[start_idx:i+1]
                    break
        
        if json_match:
            try:
                # Clean control characters that can cause JSON decode errors
                cleaned_json = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_match)
                result = json.loads(cleaned_json)
                is_ambiguous = result.get("is_ambiguous", False)
                clarification = result.get("clarification_suggestion", "") if is_ambiguous else ""
                
                logger.info(f"Ambiguous query detection: is_ambiguous={is_ambiguous}, query='{query[:50]}...'")
                return is_ambiguous, clarification
            except json.JSONDecodeError as je:
                logger.warning(f"JSON decode error: {je}, cleaned JSON preview: {cleaned_json[:200] if 'cleaned_json' in locals() else 'N/A'}, using fallback")
                return _fallback_ambiguous_detection(query, conversation_context)
        else:
            # Fallback: if JSON parsing fails, use heuristics
            logger.warning("Failed to extract JSON from LLM response for ambiguous query detection, using fallback")
            return _fallback_ambiguous_detection(query, conversation_context)
    except Exception as e:
        logger.error(f"Error in detect_ambiguous_query: {e}")
        # Fallback to heuristic detection
        return _fallback_ambiguous_detection(query, conversation_context)


def _fallback_ambiguous_detection(query: str, conversation_context: str = "") -> Tuple[bool, str]:
    """
    Fallback heuristic-based ambiguous query detection.
    Used when LLM detection fails.
    IMPORTANT: This function considers conversation context - queries are only ambiguous if there's no context.
    """
    query_lower = query.strip().lower()
    
    # FIRST: Check if query mentions a specific compliance framework
    # If it does, it's NOT ambiguous even if it's in a clarification loop
    framework_patterns = [
        r'\bgdpr\b', r'\biso\s*27001\b', r'\biso\s*27002\b', r'\biso/iec\s*27001\b', r'\biso/iec\s*27002\b',
        r'\bsoc\s*2\b', r'\bsoc2\b', r'\bsoc\s*ii\b',
        r'\bhipaa\b',
        r'\bpci\s*dss\b', r'\bpci-dss\b', r'\bpci\b',
        r'\bnist\b',
        r'\bccpa\b', r'\bcpra\b',
        r'\biso\s*22301\b', r'\biso\s*31000\b'
    ]
    
    mentions_framework = any(re.search(pattern, query_lower, re.IGNORECASE) for pattern in framework_patterns)
    
    if mentions_framework:
        # Query mentions a framework - NOT ambiguous, proceed to answer
        logger.info(f"Fallback: Query mentions framework: '{query}' - NOT ambiguous")
        return False, ""
    
    # Check if there's meaningful conversation context
    has_context = conversation_context and len(conversation_context.strip()) > 50
    
    # FIRST: Check if previous bot response was asking for clarification
    # If so, and current query is still vague, treat as ambiguous
    if has_context:
        context_lower = conversation_context.lower()
        clarification_indicators = [
            "could you please specify", "could you please provide", "please specify", "please provide",
            "what 'that thing' is", "what compliance framework", "which compliance framework",
            "what specific requirement", "what specific control", "i'd be happy to help",
            "i'd be happy to provide", "clarification"
        ]
        is_followup_to_clarification = any(indicator in context_lower for indicator in clarification_indicators)
        
        if is_followup_to_clarification:
            # Check if current query is still vague
            vague_patterns = [
                r'^(that|this|it|the thing|that thing|this thing|them|those|these)$',
                r'^(what|which|how|why|when|where)\s+(is|are|was|were|do|does|did|will|would|can|could|should|must)\s*(that|this|it|the thing|that thing)?\s*\?*$',
                r'^(tell me|explain|describe|what about|how about)\s*(that|this|it|the thing|that thing)?\s*\?*$',
            ]
            query_normalized = re.sub(r'\s+', ' ', query_lower).strip()
            for pattern in vague_patterns:
                if re.match(pattern, query_normalized, re.IGNORECASE):
                    logger.info(f"Fallback: Vague follow-up to clarification request detected: '{query}'")
                    return True, "I still need more specific information to help you. Could you please provide details such as:\n- Which compliance framework are you asking about? (GDPR, ISO 27001, SOC 2, HIPAA, PCI DSS, etc.)\n- What specific requirement, control, or regulation?\n- What aspect would you like to know more about? (implementation, requirements, certification, etc.)"
    
    # If there's context, check if it contains compliance-related terms that would make follow-ups clear
    context_has_compliance_info = False
    if has_context:
        context_lower = conversation_context.lower()
        compliance_indicators = ['gdpr', 'iso', 'soc', 'hipaa', 'pci', 'nist', 'ccpa', 'framework', 'control', 
                                'requirement', 'article', 'standard', 'compliance', 'regulation', 'audit', 
                                'certification', 'encryption', 'access', 'security', 'privacy', 'data', 'breach', 
                                'notification', 'requirement', 'control', 'standard']
        context_has_compliance_info = any(indicator in context_lower for indicator in compliance_indicators)
    
    # Patterns that indicate potentially ambiguous queries (but only if no context)
    # Note: Using more flexible patterns to catch variations and typos
    ambiguous_patterns = [
        # Very vague single phrases (e.g., "that thing", "this thing", "it")
        r'^(that|this|it|the\s+thing|that\s+thing|this\s+thing|them|those|these)$',
        # Exact matches for common ambiguous phrases
        r'^(tell\s*me\s*more|what\s*about|how\s*about|what\'?s\s*required|what\s*is\s*required|how\s*do\s*i\s*do\s*it|how\s*do\s*you\s*do\s*it|what\s*do\s*i\s*need|what\s*should\s*i\s*do|what\s*are\s*the\s*steps|how\s*can\s*i|what\s*can\s*i\s*do)$',
        # "tell me more" with optional words after (e.g., "tell me more about that thing")
        r'^tell\s*me\s*more(\s+about|\s+on|\s+regarding)?(\s+(that|this|it|the)\s+(thing|one|stuff|topic|subject))?',
        # Vague questions with pronouns/references
        r'^(tell\s*me|explain|describe|what|how|why|when|where)\s+(more|about|it|this|that|them|these|those)(\s+(thing|one|stuff|topic|subject))?',
        # Generic questions without specifics
        r'^(what|how|why|when|where)\s+(about|is|are|do|does|did|will|would|can|could|should|must)\s*(it|this|that|them|the\s+thing)?\s*\?*$',
        # "what about that thing" pattern
        r'^what\s+about\s+(that|this|it|the)\s+(thing|one|stuff|topic|subject)',
    ]
    
    # Check if query matches ambiguous patterns (case-insensitive, flexible whitespace)
    matches_ambiguous_pattern = False
    # Normalize whitespace for better matching
    query_normalized = re.sub(r'\s+', ' ', query_lower).strip()
    
    for pattern in ambiguous_patterns:
        if re.match(pattern, query_normalized, re.IGNORECASE):
            matches_ambiguous_pattern = True
            logger.info(f"Fallback: Query matches ambiguous pattern: '{query}' -> pattern: {pattern}")
            break
    
    # If query matches ambiguous pattern BUT there's context with compliance info, it's NOT ambiguous
    if matches_ambiguous_pattern:
        if has_context and context_has_compliance_info:
            # This is a follow-up query with context - NOT ambiguous
            logger.info(f"Follow-up query detected with context: '{query}' - not ambiguous")
            return False, ""
        else:
            # No context or context doesn't help - ambiguous
            logger.info(f"Fallback: Ambiguous query detected (no context): '{query}'")
            return True, "I'd be happy to help! Could you please provide more details about what you'd like to know? For example:\n- What compliance framework are you interested in? (GDPR, ISO 27001, SOC 2, etc.)\n- What specific requirement or control do you need information about?\n- Are you looking for implementation guidance or regulatory requirements?"
    
    # Check if query is very short and lacks specific terms
    if len(query_lower) < 15:
        # Check for presence of specific compliance-related terms
        specific_terms = ['gdpr', 'iso', 'soc', 'hipaa', 'pci', 'nist', 'ccpa', 'framework', 'control', 'requirement', 
                         'article', 'standard', 'compliance', 'regulation', 'audit', 'certification', 'encryption',
                         'access', 'security', 'privacy', 'data', 'breach', 'notification']
        if not any(term in query_lower for term in specific_terms):
            # If no conversation context, it's likely ambiguous
            if not has_context or not context_has_compliance_info:
                return True, "I'd be happy to help! Could you please provide more details about what you'd like to know? For example:\n- What compliance framework are you interested in? (GDPR, ISO 27001, SOC 2, etc.)\n- What specific requirement or control do you need information about?\n- Are you looking for implementation guidance or regulatory requirements?"
    
    return False, ""


@timing_decorator
def detect_query_type(query: str, conversation_context: str = "") -> Tuple[str, List[str]]:
    """
    Detect the type of compliance query and required experts using enhanced selection.
    
    Args:
        query (str): The user's query
        conversation_context (str): The conversation context to help determine query type
        
    Returns:
        Tuple[str, List[str]]: (query_type, required_experts)
        - query_type: 'ambiguous' if query needs clarification, otherwise the detected type
        - required_experts: empty list if ambiguous, otherwise list of expert types
    """
    
    # First, check if query is ambiguous
    is_ambiguous, clarification_message = detect_ambiguous_query(query, conversation_context)
    if is_ambiguous:
        # Store clarification message in a special way - we'll handle this in the route
        return 'ambiguous', []
    
    # Check for framework selection queries
    framework_keywords = ['which framework', 'what framework', 'recommend framework', 
                         'choose framework', 'select framework', 'best framework',
                         'framework recommendation', 'framework selection']
    
    if any(keyword in query.lower() for keyword in framework_keywords):
        return 'framework_selection', ['audit']
    
    # Use the enhanced expert selection system for all other queries
    required_experts = select_relevant_experts_optimized(query)
    
    # Determine query type based on selected experts
    if len(required_experts) == 1:
        if required_experts[0] == 'security':
            query_type = 'security'
        elif required_experts[0] == 'privacy':
            query_type = 'privacy'
        elif required_experts[0] == 'healthcare':
            query_type = 'healthcare'
        elif required_experts[0] == 'financial':
            query_type = 'financial'
        elif required_experts[0] == 'healthcare':
            query_type = 'healthcare'
        elif required_experts[0] == 'international':
            query_type = 'international'
        elif required_experts[0] == 'operational':
            query_type = 'operational'
        elif required_experts[0] == 'industry_specific':
            query_type = 'industry_specific'
        else:
            query_type = 'audit'
    else:
        # Multi-expert query
        query_type = 'multi_domain'
    
    # Default to audit if no specific experts identified
    if not required_experts:
        required_experts = ['audit']
        query_type = 'general'
    
    return query_type, required_experts

    

    

@timing_decorator
def process_query_optimized(query: str, context: str, conversation_context: str, conversation_history: 'ConversationHistory' = None) -> Tuple[str, float]:
    """Optimized query processing with fast response paths."""
    start_time = time.time()
    
    # Step 1: Validate input
    is_valid, error_message = validate_query_input(query)
    if not is_valid:
        end_time = time.time()
        return error_message, end_time - start_time
    
    # Step 2: Check for conversation history queries
    if detect_conversation_history_query(query):
        response = handle_conversation_history_query(conversation_history)
        end_time = time.time()
        return response, end_time - start_time
    
    # Fast Path: Check if this is a compliance-related query
    is_compliance, reason = is_compliance_related_optimized(query, conversation_context)
    
    if not is_compliance:
        # Quick non-compliance response
        response = generate_non_compliance_response(query)
        end_time = time.time()
        return response, end_time - start_time
    
    # Fast Path: Check for simple informational queries
    if detect_informational_query(query):
        logger.info("Detected informational query - using concise response")
        response = generate_concise_informational_response(query, context)
        end_time = time.time()
        return response, end_time - start_time
    
    # Check for exact cache match first - use normalized query for exact matching
    # Normalize query: lowercase, strip whitespace, remove extra spaces
    query_normalized = ' '.join(query.strip().lower().split())
    exact_query_key = f"exact_query:{hash_text(query_normalized)}"
    
    # First check: exact query match (same question = same answer, regardless of context)
    if exact_query_key in QUERY_CACHE:
        cached_response = QUERY_CACHE[exact_query_key]
        if isinstance(cached_response, str) and len(cached_response) > 0:
            logger.info(f"✅ EXACT CACHE HIT for query: '{query[:60]}...' (Key: {exact_query_key})")
            end_time = time.time()
            return cached_response, end_time - start_time
    
    # Second check: query + context match (for context-dependent responses)
    context_hash = hash_text(context[:500] if context else "")
    conv_hash = hash_text(conversation_context[:200] if conversation_context else "")
    context_query_key = f"context_query:{hash_text(query_normalized)}:{context_hash}:{conv_hash}"
    
    if context_query_key in QUERY_CACHE:
        cached_response = QUERY_CACHE[context_query_key]
        if isinstance(cached_response, str) and len(cached_response) > 0:
            logger.info(f"✅ CONTEXT CACHE HIT for query: '{query[:60]}...'")
            end_time = time.time()
            return cached_response, end_time - start_time
    
    # Fast expert selection
    required_experts = select_relevant_experts_optimized(query)
    
    # Check for partial cache matches (similar queries)
    similar_response = find_similar_cached_response(query, required_experts)
    if similar_response:
        logger.info("Using similar cached response")
        end_time = time.time()
        return similar_response, end_time - start_time
    
    # Process with optimized expert system
    expert_responses = []
    for expert in required_experts:
        response = cached_expert_response(expert, query, context, conversation_context)
        if response:
            expert_responses.append(response)
    
    # Quick aggregation for speed
    if len(expert_responses) == 1:
        final_response = expert_responses[0]
    else:
        final_response = aggregate_expert_outputs(expert_responses, query, context, conversation_context)
    
    # Cache the response with both exact query key and context-dependent key
    # Store in exact_query cache for instant retrieval on same question
    QUERY_CACHE[exact_query_key] = final_response
    
    # Also store with context for context-dependent caching
    QUERY_CACHE[context_query_key] = final_response
    
    logger.info(f"💾 Cached response for query: '{query[:60]}...' (Cache size: {len(QUERY_CACHE)})")
    
    # Save cache immediately to disk for persistence
    try:
        save_query_cache()
        logger.debug(f"✅ Cache saved successfully to {QUERY_CACHE_FILE}")
    except Exception as e:
        logger.error(f"❌ Error saving query cache: {e}")
        # Don't block on cache save errors, but log them
    
    end_time = time.time()
    return final_response, end_time - start_time

def find_similar_cached_response(query: str, experts: List[str]) -> Optional[str]:
    """Find similar cached responses for faster retrieval."""
    try:
        query_lower = query.lower()
        
        # Look for cached responses with similar expert combinations
        for cache_key, cached_data in QUERY_CACHE.items():
            if isinstance(cached_data, dict) and 'response' in cached_data:
                # Check if it's a similar query type
                if cached_data.get('experts') == experts:
                    # Simple keyword similarity check
                    cached_query = cached_data.get('original_query', '')
                    if cached_query:
                        # Count common words
                        query_words = set(query_lower.split())
                        cached_words = set(cached_query.lower().split())
                        similarity = len(query_words & cached_words) / len(query_words | cached_words)
                        
                        if similarity > 0.7:  # High similarity threshold
                            return cached_data['response']
        
        return None
        
    except Exception as e:
        logger.error(f"Error finding similar cached response: {e}")
        return None


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file with support for scanned documents."""
    try:
        logger.info(f"Extracting text from PDF: {file_path}")
        
        # First try with PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                # Validate PDF structure
                if not pdf_reader.metadata and len(pdf_reader.pages) == 0:
                    raise ValueError("Invalid PDF structure")
                
                text = ""
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"No text extracted from page {i+1}")
                    except Exception as e:
                        logger.error(f"Error extracting text from page {i+1}: {str(e)}")
                        continue
                
                if text.strip():
                    return text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        # Try pdfplumber
        try:
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        logger.error(f"Error extracting text with pdfplumber: {str(e)}")
                        continue
                
                if text.strip():
                    return text
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # If no text was extracted, try OCR using pytesseract (optional - requires poppler)
        try:
            import pytesseract
            from pdf2image import convert_from_path
            import tempfile
            
            # Check if poppler is available by trying to convert first page
            try:
                # Convert PDF to images
                with tempfile.TemporaryDirectory() as temp_dir:
                    images = convert_from_path(file_path, output_folder=temp_dir)
                    
                    if not images:
                        raise ValueError("No images generated from PDF")
                    
                    text = ""
                    for i, image in enumerate(images):
                        try:
                            # Extract text using OCR
                            page_text = pytesseract.image_to_string(image)
                            if page_text:
                                text += page_text + "\n"
                            else:
                                logger.warning(f"No text extracted from page {i+1} using OCR")
                        except Exception as e:
                            logger.error(f"Error during OCR for page {i+1}: {str(e)}")
                            continue
                    
                    if text.strip():
                        return text
            except Exception as pdf2img_error:
                error_str = str(pdf2img_error)
                if "poppler" in error_str.lower() or "Unable to get page count" in error_str:
                    # Poppler not installed - skip OCR
                    logger.warning("Poppler not available - OCR skipped. Install poppler for scanned PDF support.")
                    raise ValueError("OCR requires poppler to be installed. See INSTALL_POPPLER_WINDOWS.md for instructions.")
                raise
        except ImportError as import_error:
            logger.warning(f"OCR libraries not available: {str(import_error)}")
            ocr_error_msg = f"OCR libraries not installed: {str(import_error)}"
        except ValueError as value_error:
            # Re-raise our custom poppler error
            if "poppler" in str(value_error).lower():
                ocr_error_msg = str(value_error)
                raise
            raise
        except Exception as ocr_error:
            logger.warning(f"OCR extraction failed: {str(ocr_error)}")
            ocr_error_msg = str(ocr_error)
        
        # If all methods fail, try to repair the PDF
        try:
            import pikepdf
            with pikepdf.open(file_path) as pdf:
                # Create a temporary file for the repaired version
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    repaired_path = temp_file.name
                
                # Save repaired version
                pdf.save(repaired_path)
                
                # Try extraction again with the repaired file
                try:
                    with open(repaired_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                        
                        if text.strip():
                            return text
                finally:
                    # Clean up repaired file
                    try:
                        os.remove(repaired_path)
                    except:
                        pass
        except Exception as repair_error:
            logger.error(f"PDF repair attempt failed: {str(repair_error)}")
        
        # Provide more specific error message based on what failed
        error_details = []
        
        # Check if OCR failed due to poppler (from the exception message)
        if 'ocr_error_msg' in locals() and ocr_error_msg:
            if "poppler" in ocr_error_msg.lower() or "Unable to get page count" in ocr_error_msg:
                error_details.append("OCR processing requires poppler to be installed")
        
        error_msg = "Unable to extract text from the PDF document."
        if error_details:
            error_msg += f" {error_details[0]}."
        error_msg += " The PDF might be image-based (scanned), password-protected, or corrupted. Please ensure the PDF contains selectable text or try converting it to a text-based PDF format."
        
        raise ValueError(error_msg)
        
    except ValueError as e:
        # Re-raise our custom error messages as-is
        if "Unable to extract text" in str(e):
            raise
        # Otherwise, wrap it
        raise ValueError(f"PDF extraction failed: {str(e)}")
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        # Wrap in a user-friendly message
        error_msg = f"PDF extraction failed: {str(e)}"
        if "poppler" in str(e).lower():
            error_msg = "Unable to extract text from the PDF document. OCR processing requires poppler to be installed. The PDF might be image-based (scanned). Please ensure the PDF contains selectable text or try converting it to a text-based PDF format."
        raise ValueError(error_msg)

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    try:
        logger.info(f"Extracting text from DOCX: {file_path}")
        doc = docx.Document(file_path)
        text = ""
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                if paragraph.text:
                    text += paragraph.text + "\n"
            except Exception as e:
                logger.error(f"Error extracting text from paragraph {i+1}: {str(e)}")
                continue
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def process_uploaded_document(file_path: str) -> str:
    """Process an uploaded document and return its text content."""
    try:
        logger.info(f"Processing document: {file_path}")
        
        # Determine file type and extract text
        if file_path.lower().endswith('.pdf'):
            logger.info("Extracting text from PDF")
            text = extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            logger.info("Extracting text from DOCX")
            text = extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path}")

        if not text or not text.strip():
            raise ValueError("No text content could be extracted from the document")

        logger.info(f"Successfully extracted {len(text)} characters of text")
        
        # Process the text into segments
        segments = process_text_into_segments(text)
        logger.info(f"Created {len(segments)} text segments")
        
        # Generate embeddings for the segments
        embeddings = []
        for i, segment in enumerate(segments):
            try:
                embedding = get_embedding(segment)
                if embedding is not None:
                    embeddings.append(embedding)
                else:
                    logger.warning(f"Failed to generate embedding for segment {i}")
            except Exception as e:
                logger.error(f"Error generating embedding for segment {i}: {str(e)}")
                continue
        
        if not embeddings:
            raise ValueError("Failed to generate embeddings for any segments")
        
        logger.info(f"Generated {len(embeddings)} embeddings")
        
        # Create or update FAISS index
        try:
            embeddings_array = np.array(embeddings).astype('float32')
            dimension = embeddings_array.shape[1]
            
            # Create a new index or load existing one
            index_path = Path("faiss_indexes") / "uploaded_docs.index"
            if index_path.exists():
                logger.info("Loading existing FAISS index")
                index = faiss.read_index(str(index_path))
                # Add new vectors to existing index
                index.add(embeddings_array)
            else:
                logger.info("Creating new FAISS index")
                # Create new index
                index = faiss.IndexFlatL2(dimension)
                index.add(embeddings_array)
            
            # Save the updated index
            index_path.parent.mkdir(exist_ok=True)
            faiss.write_index(index, str(index_path))
            logger.info("FAISS index saved successfully")
            
            # Save segments and embeddings for future use
            segments_path = Path("embeddings") / "uploaded_docs_segments.json"
            embeddings_path = Path("embeddings") / "uploaded_docs_embeddings.npy"
            
            segments_path.parent.mkdir(exist_ok=True)
            embeddings_path.parent.mkdir(exist_ok=True)
            
            with open(segments_path, 'w', encoding='utf-8') as f:
                json.dump(segments, f)
            
            np.save(str(embeddings_path), embeddings_array)
            logger.info("Segments and embeddings saved successfully")
            
        except Exception as e:
            logger.error(f"Error creating/updating FAISS index: {str(e)}")
            # Continue even if index creation fails - we still have the text
        
        return text

    except Exception as e:
        logger.error(f"Error processing uploaded document: {str(e)}", exc_info=True)
        raise

def process_text_into_segments(text: str, max_segment_length: int = 1000) -> List[str]:
    """Process text into segments of appropriate length."""
    try:
        logger.info("Processing text into segments")
        # Split text into sentences (simple approach)
        sentences = text.replace('\n', ' ').split('. ')
        segments = []
        current_segment = ""
        
        for sentence in sentences:
            if len(current_segment) + len(sentence) < max_segment_length:
                current_segment += sentence + ". "
            else:
                if current_segment:
                    segments.append(current_segment.strip())
                current_segment = sentence + ". "
        
        if current_segment:
            segments.append(current_segment.strip())
        
        logger.info(f"Created {len(segments)} segments")
        return segments
    except Exception as e:
        logger.error(f"Error processing text into segments: {str(e)}")
        raise

def generate_terms_and_conditions(framework: str, format: str = "txt") -> str:
    """
    Generate Terms & Conditions document based on a specific compliance framework.
    """
    try:
        # Get framework requirements
        requirements = get_framework_requirements(framework)
        
        # Generate sections
        sections = []
        
        # Introduction
        sections.append(f"Terms and Conditions\n")
        sections.append(f"Last Updated: {datetime.now().strftime('%Y-%m-%d')}\n")
        sections.append(f"These Terms and Conditions govern your use of our services and are designed to comply with {framework} requirements.\n")
        
        # Generate content for each requirement using AI
        for req in requirements:
            section_title = req['title']
            prompt = (
                f"Generate a detailed section for {framework} compliant Terms and Conditions about {req['title']}.\n"
                f"Requirement: {req['description']}\n"
                "The section should:\n"
                "1. Be specific and actionable\n"
                "2. Include all necessary legal requirements\n"
                "3. Use clear, professional language\n"
                "4. Be comprehensive but concise\n"
                "5. Include practical implementation details\n\n"
                "Generate the section content:"
            )
            
            section_content = rate_limited_generate_content(prompt)
            sections.append(f"\n{section_title}\n")
            sections.append(f"{section_content}\n")
        
        # Add standard sections
        standard_sections = {
            "Acceptance of Terms": "By accessing or using our services, you agree to be bound by these Terms and Conditions.",
            "Service Description": "We provide compliance and security services designed to help organizations meet regulatory requirements.",
            "User Obligations": "Users must comply with all applicable laws and regulations while using our services.",
            "Intellectual Property": "All content and materials provided through our services are protected by intellectual property rights.",
            "Limitation of Liability": "We are not liable for any indirect, incidental, special, consequential, or punitive damages.",
            "Termination": "We reserve the right to terminate or suspend access to our services for violations of these terms.",
            "Governing Law": f"These terms are governed by applicable laws and {framework} requirements.",
            "Contact Information": "For questions about these Terms and Conditions, please contact us at:\n\n[Company Name]\n[Address]\n[Email]\n[Phone]"
        }
        
        for title, content in standard_sections.items():
            sections.append(f"\n{title}\n")
            sections.append(f"{content}\n")
        
        # Combine all sections
        terms_text = "\n".join(sections)
        
        # Convert to requested format
        if format.lower() == "pdf":
            # Convert to PDF
            doc = Document()
            for section in sections:
                doc.add_paragraph(section)
            pdf_path = f"generated_terms_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            doc.save(pdf_path)
            return f"Terms and Conditions generated and saved as {pdf_path}"
            
        elif format.lower() == "docx":
            # Convert to DOCX
            doc = Document()
            for section in sections:
                doc.add_paragraph(section)
            docx_path = f"generated_terms_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(docx_path)
            return f"Terms and Conditions generated and saved as {docx_path}"
            
        else:
            return terms_text
            
    except Exception as e:
        logger.error(f"Error generating Terms and Conditions: {str(e)}")
        return f"Error generating Terms and Conditions: {str(e)}"


# Dynamic learning system for query classification
CLASSIFICATION_FEEDBACK_FILE = os.path.join(CACHE_DIR, "classification_feedback.json")
classification_feedback = {}

# Load existing feedback
if os.path.exists(CLASSIFICATION_FEEDBACK_FILE):
    try:
        with open(CLASSIFICATION_FEEDBACK_FILE, 'r') as f:
            classification_feedback = json.load(f)
        logger.info(f"Loaded {len(classification_feedback)} classification feedback entries")
    except Exception as e:
        logger.error(f"Error loading classification feedback: {e}")

def save_classification_feedback():
    """Save classification feedback to disk"""
    try:
        with open(CLASSIFICATION_FEEDBACK_FILE, 'w') as f:
            json.dump(classification_feedback, f, indent=2)
        logger.info(f"Saved {len(classification_feedback)} classification feedback entries")
    except Exception as e:
        logger.error(f"Error saving classification feedback: {e}")

def learn_from_user_interaction(query: str, was_helpful: bool, actual_classification: bool):
    """Learn from user feedback to improve future classifications"""
    query_hash = hash_text(query)
    
    feedback_entry = {
        'query': query,
        'was_helpful': was_helpful,
        'actual_classification': actual_classification,
        'timestamp': datetime.now().isoformat(),
        'embedding': get_embedding(query).tolist() if get_embedding(query) is not None else None
    }
    
    classification_feedback[query_hash] = feedback_entry
    save_classification_feedback()

    

def analyze_document_intent(query: str, conversation_context: str = "", has_uploaded_doc: bool = False) -> Dict[str, Any]:
    """
    Intelligently analyze user intent regarding document operations using AI.
    Returns classification of what the user wants to do with documents.
    """
    try:
        prompt = f"""
Analyze this user query to understand their intent. Be VERY precise in classification.

User Query: "{query}"
Conversation Context: "{conversation_context}"
Has Uploaded Document: {has_uploaded_doc}

CLASSIFICATION RULES:
1. GENERAL_COMPLIANCE - Simple informational questions about compliance concepts, frameworks, regulations
   Examples: "what is gdpr", "explain hipaa", "what are gdpr requirements", "tell me about ccpa"

2. ANALYZE_UPLOADED - Only when user explicitly asks about an uploaded document
   Examples: "analyze my document", "review my privacy policy", "check my uploaded file"

3. GENERATE_NEW - Only when user EXPLICITLY asks to CREATE/GENERATE a new document
   Examples: "create a privacy policy", "generate gdpr document", "make me a terms document"
   NOT: "what is gdpr", "explain privacy policy", "gdpr requirements"

4. COMPARE_COMPLIANCE - User wants to compare their document against standards
   Examples: "compare my document to gdpr", "how compliant is my policy"

5. GET_IMPROVEMENT_SUGGESTIONS - User wants suggestions (no document generation)
   Examples: "how can I improve", "what suggestions do you have", "how to make it better"

6. IMPROVE_DOCUMENT - User explicitly wants an improved version generated
   Examples: "generate improved version", "create better document", "make improved policy"

7. NON_DOCUMENT - Not related to compliance or documents

CRITICAL: 
- Simple questions about compliance topics = GENERAL_COMPLIANCE
- Questions about regulations/frameworks = GENERAL_COMPLIANCE  
- Only classify as GENERATE_NEW if user explicitly asks to CREATE/GENERATE/MAKE a document
- "what is X" questions are ALWAYS GENERAL_COMPLIANCE

Current Query Analysis:
- Is this asking WHAT something is? → GENERAL_COMPLIANCE
- Is this asking to CREATE/GENERATE? → GENERATE_NEW
- Is this about an uploaded document? → ANALYZE_UPLOADED

Respond in JSON format:
{{
  "intent": "CATEGORY_NAME",
  "document_type": "unknown",
  "framework": "general",
  "urgency": "low",
  "confidence": 0.0,
  "reasoning": "brief explanation"
}}
"""

        response = rate_limited_generate_content(prompt, temperature=0.1)
        
        try:
            # Try to parse JSON response
            import json
            intent_data = json.loads(response)
            
            # Validate required fields
            required_fields = ["intent", "document_type", "framework", "urgency", "confidence", "reasoning"]
            if all(field in intent_data for field in required_fields):
                return intent_data
        except:
            pass
        
        # Fallback analysis if AI parsing fails
        return analyze_document_intent_fallback(query, has_uploaded_doc)
        
    except Exception as e:
        logger.error(f"Error in document intent analysis: {e}")
        return analyze_document_intent_fallback(query, has_uploaded_doc)

def analyze_document_intent_fallback(query: str, has_uploaded_doc: bool) -> Dict[str, Any]:
    """Fallback document intent analysis using semantic similarity."""
    try:
        query_lower = query.lower()
        query_embedding = get_embedding(query)
        
        # First check for informational questions - these should NEVER be document generation
        informational_patterns = [
            "what is", "what are", "explain", "tell me about", "describe", 
            "define", "definition of", "meaning of", "what does", "how does"
        ]
        
        if any(pattern in query_lower for pattern in informational_patterns):
            # Check if it's about a framework
            framework = "general"
            frameworks = ["gdpr", "ccpa", "hipaa", "iso27001", "soc2", "nist", "pci"]
            for fw in frameworks:
                if fw in query_lower:
                    framework = fw.upper()
                    break
            
            return {
                "intent": "GENERAL_COMPLIANCE",
                "document_type": "unknown",
                "framework": framework,
                "urgency": "low",
                "confidence": 0.95,
                "reasoning": "Informational query about compliance topic"
            }
        
        # Define intent patterns with embeddings - updated with better patterns
        intent_patterns = {
            "ANALYZE_UPLOADED": [
                "analyze my document for compliance issues",
                "check if my privacy policy is compliant",
                "review my terms and conditions",
                "is my document following regulations"
            ],
            "GENERATE_NEW": [
                "create a new privacy policy for me",
                "generate compliant terms and conditions", 
                "make me a GDPR compliant document",
                "draft a privacy policy according to regulations"
            ],
            "COMPARE_COMPLIANCE": [
                "compare my document against GDPR requirements",
                "check compliance with specific framework",
                "how does my policy measure against standards"
            ],
            "GET_IMPROVEMENT_SUGGESTIONS": [
                "how can I make it better",
                "how to improve my document",
                "what should I fix in my policy",
                "give me suggestions to improve",
                "tell me how to make it better",
                "what improvements do you recommend"
            ],
            "IMPROVE_DOCUMENT": [
                "generate an improved version of my document",
                "create a better version of my policy",
                "make me an improved compliant document",
                "fix my document and generate new version"
            ]
        }
        
        best_intent = "GENERAL_COMPLIANCE"
        best_score = 0.0
        
        # Check for explicit generation keywords only if they exist
        generation_keywords = ["create", "generate", "make me", "draft", "produce", "build"]
        document_keywords = ["document", "policy", "terms", "agreement", "contract"]
        
        has_generation_intent = any(gen_word in query_lower for gen_word in generation_keywords)
        has_document_context = any(doc_word in query_lower for doc_word in document_keywords)
        
        # Only consider generation if BOTH generation keywords AND document context exist
        if has_generation_intent and has_document_context:
            # Check for specific suggestion keywords first
            suggestion_keywords = ["how can i", "how to", "what should i", "suggestions", "recommend", "advice", "tell me how"]
            if any(keyword in query_lower for keyword in suggestion_keywords):
                best_intent = "GET_IMPROVEMENT_SUGGESTIONS"
                best_score = 0.9
            else:
                # Use embedding similarity for generation intents
                if query_embedding is not None:
                    for intent, patterns in intent_patterns.items():
                        for pattern in patterns:
                            pattern_embedding = get_embedding(pattern)
                            if pattern_embedding is not None:
                                similarity = cosine_similarity([query_embedding], [pattern_embedding])[0][0]
                                if similarity > best_score and similarity > 0.7:  # Higher threshold for generation
                                    best_score = similarity
                                    best_intent = intent
        
        # Determine document type
        doc_type = "unknown"
        if any(term in query_lower for term in ["privacy policy", "privacy"]):
            doc_type = "privacy_policy"
        elif any(term in query_lower for term in ["terms", "conditions", "terms and conditions"]):
            doc_type = "terms_conditions"
        
        # Determine framework
        framework = "general"
        frameworks = ["gdpr", "ccpa", "hipaa", "iso27001", "soc2", "nist", "pci"]
        for fw in frameworks:
            if fw in query_lower:
                framework = fw.upper()
                break
        
        return {
            "intent": best_intent,
            "document_type": doc_type,
            "framework": framework,
            "urgency": "medium",
            "confidence": best_score,
            "reasoning": f"Semantic analysis with {best_score:.2f} confidence"
        }
        
    except Exception as e:
        logger.error(f"Error in fallback intent analysis: {e}")
        return {
            "intent": "GENERAL_COMPLIANCE",
            "document_type": "unknown", 
            "framework": "general",
            "urgency": "low",
            "confidence": 0.0,
            "reasoning": "Fallback analysis"
        }

def generate_comprehensive_document_analysis(document_text: str, framework: str, document_type: str) -> str:
    """Generate a comprehensive analysis of an uploaded document against compliance requirements."""
    try:
        prompt = f"""
Conduct a comprehensive compliance analysis of this {document_type} document against {framework} requirements.

Document Content (first 3000 characters):
{document_text[:3000]}

Provide a detailed analysis including:

1. **Executive Summary**
   - Overall compliance status (Compliant/Partially Compliant/Non-Compliant)
   - Compliance score (percentage)
   - Key findings summary

2. **Detailed Requirements Analysis**
   - List specific {framework} requirements
   - For each requirement: Met/Not Met/Partially Met
   - Evidence from the document
   - Missing elements

3. **Critical Issues** (if any)
   - High-priority compliance gaps
   - Legal risks
   - Immediate action items

4. **Recommendations**
   - Specific improvements needed
   - Additional clauses to add
   - Language modifications

5. **Implementation Roadmap**
   - Priority 1: Critical fixes
   - Priority 2: Important improvements  
   - Priority 3: Best practice enhancements

6. **Next Steps**
   - Immediate actions required
   - Timeline recommendations
   - Follow-up considerations

Format the response with clear headings and bullet points for easy reading.
"""

        analysis = rate_limited_generate_content(prompt, temperature=0.2)
        
        if not analysis or len(analysis.strip()) < 100:
            return f"I've analyzed your {document_type} document against {framework} requirements. The document appears to have some compliance gaps that need attention. Would you like me to generate a fully compliant version for you?"
        
        return analysis
        
    except Exception as e:
        logger.error(f"Error generating document analysis: {e}")
        return f"I've reviewed your {document_type} document. There appear to be some areas that could be improved for {framework} compliance. Would you like me to create a fully compliant version for you?"

def generate_intelligent_compliant_document(document_type: str, framework: str, organization_context: str = "") -> str:
    """Generate a fully compliant document based on the framework and user context."""
    try:
        prompt = f"""
Generate a comprehensive, legally compliant {document_type} document that fully meets {framework} requirements.

Organization Context: {organization_context if organization_context else "General business organization"}
Document Type: {document_type}
Compliance Framework: {framework}

IMPORTANT FORMATTING INSTRUCTIONS:
- Generate ONLY clean, plain text content suitable for a professional legal document
- DO NOT use HTML tags, markdown syntax, or special formatting codes
- Use simple numbered sections (1., 2., 3., etc.) for main headings
- Use lettered subsections (a., b., c., etc.) for subsections when needed
- Use simple dashes (-) for bullet points
- Separate sections with double line breaks
- Use [PLACEHOLDERS] in brackets for items that need customization

The document should include:

1. **Complete Legal Framework Coverage**
   - All mandatory {framework} requirements
   - Appropriate legal language
   - Jurisdiction-specific considerations

2. **Professional Structure**
   - Clear section numbering and headings
   - Logical flow and organization
   - Proper paragraph structure

3. **Comprehensive Content**
   - All required disclosures
   - User rights and obligations
   - Data handling procedures (if applicable)
   - Contact information sections
   - Legal compliance statements

4. **Actionable Implementation**
   - Clear, understandable language
   - Specific procedures and processes
   - Compliance verification methods

5. **Future-Proof Elements**
   - Regulatory change adaptability
   - Update mechanisms
   - Version control considerations

Generate a complete, ready-to-use document that an organization can immediately implement.
Make it professional, legally sound, and fully compliant with {framework} standards.

Use placeholders like [COMPANY NAME], [CONTACT EMAIL], [ADDRESS], [DATE], etc. where customization is needed.

Start with the document title and generate the complete content in plain text format:
"""

        document_content = rate_limited_generate_content(prompt, temperature=0.2)
        
        if not document_content or len(document_content.strip()) < 500:
            # Fallback to template-based generation
            return generate_template_document(document_type, framework)
        
        # Clean any HTML content that might have been generated
        cleaned_content = clean_html_content(document_content)
        
        return cleaned_content
        
    except Exception as e:
        logger.error(f"Error generating compliant document: {e}")
        return generate_template_document(document_type, framework)

def clean_html_content(content: str) -> str:
    """Remove HTML tags and clean up content for DOCX generation."""
    try:
        import re
        
        # Remove HTML tags but preserve content
        content = re.sub(r'<[^>]+>', '', content)
        
        # Convert HTML entities to their text equivalents
        html_entities = {
            '&amp;': '&',
            '&lt;': '<',
            '&gt;': '>',
            '&quot;': '"',
            '&#39;': "'",
            '&apos;': "'",
            '&nbsp;': ' ',
            '&copy;': '©',
            '&reg;': '®',
            '&trade;': '™'
        }
        
        for entity, replacement in html_entities.items():
            content = content.replace(entity, replacement)
        
        # Remove excessive whitespace while preserving document structure
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)  # Max 2 consecutive newlines
        content = re.sub(r'[ \t]+', ' ', content)  # Multiple spaces/tabs to single space
        content = re.sub(r' +\n', '\n', content)  # Remove trailing spaces
        content = re.sub(r'\n +', '\n', content)  # Remove leading spaces after newlines
        
        # Ensure proper line breaks and clean up
        content = content.strip()
        
        return content
        
    except Exception as e:
        logger.error(f"Error cleaning HTML content: {e}")
        return content

def generate_template_document(document_type: str, framework: str) -> str:
    """Generate a template document as fallback."""
    templates = {
        "privacy_policy": {
            "GDPR": """
PRIVACY POLICY

Last Updated: [DATE]

1. INTRODUCTION
[COMPANY NAME] is committed to protecting your personal data in accordance with the General Data Protection Regulation (GDPR).

2. DATA CONTROLLER
[COMPANY NAME]
[ADDRESS]
[CONTACT EMAIL]

3. PERSONAL DATA WE COLLECT
We collect and process the following categories of personal data:
- Identity data
- Contact data
- Technical data
- Usage data

4. LEGAL BASIS FOR PROCESSING
We process your personal data based on:
- Consent
- Contract performance
- Legal obligation
- Legitimate interests

5. YOUR RIGHTS UNDER GDPR
You have the right to:
- Access your personal data
- Rectify inaccurate data
- Erase your data
- Restrict processing
- Data portability
- Object to processing

6. DATA RETENTION
We retain your personal data only for as long as necessary for the purposes outlined in this policy.

7. DATA SECURITY
We implement appropriate technical and organizational measures to protect your personal data.

8. INTERNATIONAL TRANSFERS
Any international transfers of your personal data are conducted with appropriate safeguards in place.

9. CONTACT US
For any questions about this privacy policy or your personal data, please contact us at [CONTACT EMAIL].
            """,
            "CCPA": """
PRIVACY POLICY

Last Updated: [DATE]

CALIFORNIA CONSUMER PRIVACY ACT (CCPA) NOTICE

1. CATEGORIES OF PERSONAL INFORMATION
We collect the following categories of personal information:
- Identifiers
- Commercial information
- Internet activity
- Geolocation data

2. SOURCES OF PERSONAL INFORMATION
We collect personal information from:
- Directly from you
- Automatically through our services
- From third parties

3. BUSINESS PURPOSES
We use personal information for:
- Providing services
- Business operations
- Legal compliance

4. YOUR CALIFORNIA RIGHTS
Under CCPA, you have the right to:
- Know what personal information we collect
- Delete personal information
- Opt-out of sale of personal information
- Non-discrimination for exercising rights

5. HOW TO EXERCISE YOUR RIGHTS
To exercise your rights, contact us at [CONTACT EMAIL] or [PHONE NUMBER].

6. CONTACT INFORMATION
[COMPANY NAME]
[ADDRESS]
[CONTACT EMAIL]
            """
        },
        "terms_conditions": {
            "GDPR": """
TERMS AND CONDITIONS

Last Updated: [DATE]

1. ACCEPTANCE OF TERMS
By using our services, you agree to these terms and our Privacy Policy.

2. SERVICES DESCRIPTION
[COMPANY NAME] provides [DESCRIPTION OF SERVICES].

3. USER OBLIGATIONS
You agree to:
- Use services lawfully
- Provide accurate information
- Respect intellectual property rights

4. DATA PROTECTION
We process your personal data in accordance with GDPR. See our Privacy Policy for details.

5. LIMITATION OF LIABILITY
Our liability is limited to the maximum extent permitted by law.

6. GOVERNING LAW
These terms are governed by [JURISDICTION] law.

7. CONTACT INFORMATION
[COMPANY NAME]
[ADDRESS]
[CONTACT EMAIL]
            """
        }
    }
    
    if document_type in templates and framework in templates[document_type]:
        return templates[document_type][framework]
    else:
        return f"Template {document_type} document for {framework} compliance - Please customize with your specific requirements."

def create_docx_with_download_link(content: str, document_type: str, framework: str, user_id: str) -> Tuple[str, str]:
    """
    Create a DOCX file from content and return the file path and download URL.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        import os
        from datetime import datetime
        import re
        
        # Clean content first
        content = clean_html_content(content)
        
        # Create document
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Create a style for headings if it doesn't exist
        try:
            heading_style = styles['Heading 1']
        except KeyError:
            heading_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
        
        # Add title
        title_para = doc.add_heading(
            f"{document_type.replace('_', ' ').title()} - {framework} Compliant", 
            level=0
        )
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add separator line
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Empty line
        
        # Process content line by line for better formatting
        lines = content.split('\n')
        current_list = None
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                # Empty line - add paragraph break
                doc.add_paragraph()
                i += 1
                continue
            
            # Check if it's a main heading (looks like "1. TITLE" or "TITLE" in caps)
            if (re.match(r'^\d+\.\s+[A-Z][A-Z\s]+$', line) or 
                (line.isupper() and len(line) < 80 and not line.startswith('-'))):
                
                # Add as heading
                heading = doc.add_heading(line, level=1)
                current_list = None
                
            # Check if it's a numbered subsection (like "1.1", "a.", etc.)
            elif re.match(r'^[a-z]\.|^\d+\.\d+', line):
                para = doc.add_paragraph(line)
                para.style = 'List Number'
                current_list = None
                
            # Check if it's a bullet point
            elif line.startswith('-') or line.startswith('•'):
                bullet_text = line[1:].strip()  # Remove bullet character
                if current_list is None:
                    current_list = doc.add_paragraph(bullet_text, style='List Bullet')
                else:
                    doc.add_paragraph(bullet_text, style='List Bullet')
                    
            # Regular paragraph
            else:
                para = doc.add_paragraph(line)
                current_list = None
            
            i += 1
        
        # Add footer note
        doc.add_paragraph()
        footer_para = doc.add_paragraph("Note: Please customize all placeholder values (shown in [BRACKETS]) with your specific organizational information before implementing this document.")
        footer_para.style.font.size = Pt(10)
        footer_para.style.font.italic = True
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{document_type}_{framework}_{user_id}_{timestamp}.docx"
        
        # Ensure downloads directory exists
        downloads_dir = Path("downloads")
        downloads_dir.mkdir(exist_ok=True)
        
        file_path = downloads_dir / filename
        
        # Save document
        doc.save(str(file_path))
        
        # Create download URL
        download_url = f"/api/compliance/download/{filename}"
        
        logger.info(f"Created DOCX file: {file_path}")
        
        return str(file_path), download_url
        
    except Exception as e:
        logger.error(f"Error creating DOCX file: {e}")
        
        # Fallback: create a simple text file
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{document_type}_{framework}_{user_id}_{timestamp}.txt"
            downloads_dir = Path("downloads")
            downloads_dir.mkdir(exist_ok=True)
            file_path = downloads_dir / filename
            
            # Clean content for text file
            clean_content = clean_html_content(content)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(f"{document_type.replace('_', ' ').title()} - {framework} Compliant\n")
                f.write(f"Generated on: {datetime.now().strftime('%B %d, %Y')}\n")
                f.write("=" * 60 + "\n\n")
                f.write(clean_content)
                f.write("\n\n" + "=" * 60 + "\n")
                f.write("Note: Please customize all placeholder values (shown in [BRACKETS]) with your specific organizational information.")
            
            download_url = f"/api/compliance/download/{filename}"
            return str(file_path), download_url
            
        except Exception as e2:
            logger.error(f"Error creating fallback text file: {e2}")
            return None, None

def format_document_response_with_download(content: str, download_url: str, document_type: str, framework: str) -> str:
    """Format the response with document content and download link."""
    
    response = f"""## ✅ {document_type.replace('_', ' ').title()} Generated Successfully!

I've created a comprehensive {framework}-compliant {document_type.replace('_', ' ')} document for you.

### 📋 Document Preview:
{content[:800]}{"..." if len(content) > 800 else ""}

### 📥 Download Your Document
**[Click here to download your {document_type.replace('_', ' ')} document]({download_url})**

The document includes:
- ✅ Full {framework} compliance requirements
- ✅ Professional legal language
- ✅ Ready-to-use format with placeholders
- ✅ Comprehensive coverage of all necessary sections

### 🔧 Customization Required:
Please customize the document by replacing placeholders like:
- [COMPANY NAME] - Your organization name
- [CONTACT EMAIL] - Your contact email
- [ADDRESS] - Your business address
- [DATE] - Current date

### 📋 Next Steps:
1. Download the document using the link above
2. Review and customize the placeholders
3. Have your legal team review if needed
4. Implement in your organization

Would you like me to create any additional compliance documents or help you with specific sections?"""

    return response

def generate_document_improvement_suggestions(document_text: str, framework: str, document_type: str) -> str:
    """Generate specific improvement suggestions for an uploaded document without creating a new document."""
    try:
        prompt = f"""
Analyze this {document_type} document and provide specific, actionable improvement suggestions for {framework} compliance.

Document Content (first 3000 characters):
{document_text[:3000]}

Provide improvement suggestions in the following format:

## 📋 Document Improvement Recommendations

### 🔍 Quick Assessment
- **Current Compliance Level**: [Estimate percentage]
- **Framework**: {framework}
- **Priority Issues**: [Number] critical issues found

### 🚨 Critical Issues to Fix Immediately

1. **[Issue Title]**
   - **Problem**: What's missing or incorrect
   - **Impact**: Why this matters for compliance
   - **Fix**: Specific action to take
   - **Example Language**: "Suggested text to add..."

2. **[Next Issue]**
   - [Continue pattern...]

### ⚠️ Important Improvements

1. **[Issue Title]**
   - **Current**: What the document currently says
   - **Improve To**: What it should say instead
   - **Reason**: Why this change is needed

### ✨ Best Practice Enhancements

1. **[Enhancement Title]**
   - **Suggestion**: What to add or improve
   - **Benefit**: How this helps compliance

### 📝 Specific Language Suggestions

**Section: [Name]**
- **Current**: "[Existing text]"
- **Suggested**: "[Improved text]"
- **Why**: [Explanation]

### 🎯 Next Steps Priority Order

1. **Immediate (This Week)**:
   - [Action items]

2. **Short-term (This Month)**:
   - [Action items]

3. **Ongoing (Regular Review)**:
   - [Action items]

### 📚 Additional Resources
- Consider reviewing [specific standards/guidelines]
- Regular compliance audits recommended

Focus on providing specific, actionable advice rather than generating new content.
"""

        suggestions = rate_limited_generate_content(prompt, temperature=0.2)
        
        if not suggestions or len(suggestions.strip()) < 100:
            return f"""## 📋 Document Improvement Recommendations

I've reviewed your {document_type.replace('_', ' ')} document against {framework} requirements.

### 🔍 Quick Assessment
- **Framework**: {framework}
- **Review Status**: Completed

### 🚨 Key Areas for Improvement

1. **Legal Language Precision**
   - **Issue**: Some sections may need more specific legal language
   - **Action**: Review with legal counsel for framework-specific requirements

2. **Compliance Completeness**
   - **Issue**: Ensure all {framework} requirements are explicitly addressed
   - **Action**: Cross-reference against official {framework} checklist

3. **User Rights Clarity**
   - **Issue**: User rights section may need more detail
   - **Action**: Add specific procedures for users to exercise their rights

### 📝 General Recommendations
- Consider regular compliance reviews
- Ensure all stakeholders review the updated document
- Implement a document version control system

Would you like me to analyze any specific section in more detail?"""
        
        return suggestions
        
    except Exception as e:
        logger.error(f"Error generating improvement suggestions: {e}")
        return f"""## 📋 Document Improvement Recommendations

I've reviewed your {document_type.replace('_', ' ')} document for {framework} compliance.

### 🔍 General Improvement Areas

1. **Framework Alignment**: Ensure all {framework} requirements are explicitly addressed
2. **Legal Language**: Review language for clarity and legal precision
3. **User Rights**: Clearly outline user rights and how to exercise them
4. **Data Handling**: Provide detailed information about data collection and processing
5. **Contact Information**: Ensure all contact details are current and accessible

### 📝 Next Steps
1. Review each section against the {framework} checklist
2. Consider legal review of updated language
3. Regular compliance audits and updates

Would you like me to focus on any specific aspect of your document?"""

# Add in-memory embedding cache for better performance
EMBEDDING_MEMORY_CACHE = {}
EMBEDDING_CACHE_MAX_SIZE = 1000

@timing_decorator
def get_embedding_optimized(text: str) -> np.ndarray:
    """Get embeddings with aggressive in-memory caching."""
    # Normalize text for consistent caching
    text_normalized = text.strip().lower()
    text_hash = hash_text(text_normalized)
    
    # Check in-memory cache first
    if text_hash in EMBEDDING_MEMORY_CACHE:
        return EMBEDDING_MEMORY_CACHE[text_hash]
    
    # Truncate very long text for performance
    if len(text.split()) > 500:  # Reduced from 8000 for speed
        text = " ".join(text.split()[:500])
    
    try:
        embeddings = embedding_model.encode([text], convert_to_numpy=True, show_progress_bar=False)
        result = embeddings[0]
        
        # Cache in memory
        if len(EMBEDDING_MEMORY_CACHE) >= EMBEDDING_CACHE_MAX_SIZE:
            # Remove oldest entries (simple FIFO)
            oldest_key = next(iter(EMBEDDING_MEMORY_CACHE))
            del EMBEDDING_MEMORY_CACHE[oldest_key]
        
        EMBEDDING_MEMORY_CACHE[text_hash] = result
        return result
        
    except Exception as e:
        logger.error(f"Error generating embedding: {e}")
        return None

def detect_informational_query(query: str) -> bool:
    """Detect if this is a simple informational query that needs a concise response."""
    query_lower = query.lower()
    
    # Simple informational patterns
    informational_patterns = [
        "what is", "what are", "tell me about", "explain", "define",
        "definition of", "meaning of", "overview of", "introduction to"
    ]
    
    # Framework/standard specific patterns
    framework_patterns = [
        "iso 27001", "gdpr", "ccpa", "hipaa", "soc 2", "nist",
        "pci dss", "sox", "compliance framework"
    ]
    
    # Check if it's a simple "what is [framework]" type query
    is_informational = any(pattern in query_lower for pattern in informational_patterns)
    mentions_framework = any(pattern in query_lower for pattern in framework_patterns)
    
    # Look for list requests (top 10, main, key, primary)
    is_list_request = any(word in query_lower for word in [
        "top", "main", "key", "primary", "important", "rules", "requirements",
        "principles", "controls", "steps"
    ])
    
    return is_informational and mentions_framework and is_list_request

def generate_concise_informational_response(query: str, context: str = "") -> str:
    """Generate a concise, informational response for simple queries with evidence citations when context is available."""
    try:
        # Extract the framework/topic
        query_lower = query.lower()
        framework = "ISO 27001"  # Default
        
        if "gdpr" in query_lower:
            framework = "GDPR"
        elif "ccpa" in query_lower:
            framework = "CCPA"
        elif "hipaa" in query_lower:
            framework = "HIPAA"
        elif "soc 2" in query_lower or "soc2" in query_lower:
            framework = "SOC 2"
        elif "nist" in query_lower:
            framework = "NIST"
        elif "pci" in query_lower:
            framework = "PCI DSS"
        
        # Check if we have substantial framework context
        has_framework_docs = context and len(context.strip()) > 200
        
        # Check if user wants concise response
        is_concise = detect_concise_request(query)
        max_tokens = 512 if is_concise else 800
        
        if is_concise:
            # Generate a very brief response
            if has_framework_docs:
                prompt = f"""
Provide a CONCISE answer to: "{query}"

FRAMEWORK DOCUMENTS:
{context[:1000]}

CRITICAL: For each point, cite evidence from the documents using this format: (Evidence: <span style="color:#008000">"exact quote"</span> - <span style="color:#008000">Source/Control ID</span>)

Format as:
**{framework} Compliance - Key Steps:**
• Step 1 (Evidence: <span style="color:#008000">"quote"</span> - <span style="color:#008000">Article 17 GDPR</span>)
• Step 2 (Evidence: <span style="color:#008000">"quote"</span> - <span style="color:#008000">ISO 27001 A.9.2.1</span>)
• Step 3 (Evidence: <span style="color:#008000">"quote"</span> - <span style="color:#008000">Source</span>)

**MANDATORY HIGHLIGHTING:**
- All evidence quotes from framework documents MUST be highlighted in green: <span style="color:#008000">"quote"</span>
- ALL control IDs, article numbers, and framework references MUST be highlighted in green: <span style="color:#008000">Article 17 GDPR</span>, <span style="color:#008000">ISO 27001 A.9.2.1</span>

Keep under 100 words total. Focus only on essential actions with evidence citations.
"""
            else:
                prompt = f"""
Provide a CONCISE answer to: "{query}"

Format as:
**{framework} Compliance - Key Steps:**
• Step 1
• Step 2  
• Step 3
• Step 4
• Step 5

Keep under 100 words total. Focus only on essential actions.
"""
        else:
            # Generate a focused, concise response with evidence citations
            if has_framework_docs:
                prompt = f"""
Provide a clear, concise response to this query: "{query}"

FRAMEWORK DOCUMENTS:
{context}

CRITICAL INSTRUCTIONS - EVIDENCE-BASED RESPONSE REQUIREMENTS:

For EVERY requirement, rule, or claim you make, you MUST:
1. Quote the EXACT text from the framework documents above that supports your statement
2. Use this format: [Your statement] (Evidence: <span style="color:#008000">"exact quote from documents"</span> - <span style="color:#008000">Source/Section</span>)
3. If specific details are NOT in the provided documents, state: "Based on <span style=\"color:#008000\">{framework} Control/Article</span> standards" without warnings
4. **MANDATORY HIGHLIGHTING:**
   - All evidence quotes from framework documents MUST be highlighted in green: <span style="color:#008000">"quote"</span>
   - ALL control IDs, article numbers, and framework references MUST be highlighted in green: <span style="color:#008000">Article 17 GDPR</span>, <span style="color:#008000">ISO 27001 A.9.2.1</span>, <span style="color:#008000">HIPAA §164.312</span>
   - Example: Data subjects have the right to erasure (<span style="color:#008000">Article 17 GDPR</span>).

Structure your response as follows:
1. **Brief Definition** (2-3 sentences about what {framework} is)
2. **Top 10 Key Points/Rules/Requirements** (numbered list with brief explanations and evidence citations)
3. **Quick Implementation Tip** (1-2 sentences)

EXAMPLE FORMAT:
Access Control: Implement technical policies and procedures that allow only authorized personnel to access ePHI (Evidence: "Implement technical policies and procedures that allow only authorized personnel to access ePHI" - HIPAA Security Rule §164.312(a)(1))

Keep it informative but concise - aim for 300-500 words total.
Use clear, professional language that's accessible to both beginners and experts.
Focus on practical, actionable information with mandatory evidence citations.
"""
            else:
                prompt = f"""
Provide a clear, concise response to this query: "{query}"

Structure your response as follows:
1. **Brief Definition** (2-3 sentences about what {framework} is)
2. **Top 10 Key Points/Rules/Requirements** (numbered list with brief explanations)
3. **Quick Implementation Tip** (1-2 sentences)

Keep it informative but concise - aim for 300-500 words total.
Use clear, professional language that's accessible to both beginners and experts.
Focus on practical, actionable information.
"""
        
        return rate_limited_generate_content_optimized(prompt, max_tokens=max_tokens)
        
    except Exception as e:
        logger.error(f"Error generating concise response: {e}")
        return "I'd be happy to help with information about compliance frameworks. Could you please rephrase your question?"

def validate_query_input(query: str) -> Tuple[bool, str]:
    """
    Validate user query input to ensure it's meaningful and processable.
    Returns (is_valid, error_message)
    """
    if not query:
        return False, "Please enter a question or request."
    
    # Strip whitespace
    query_stripped = query.strip()
    
    if not query_stripped:
        return False, "Please enter a question or request."
    
    # Check minimum length
    if len(query_stripped) < 3:
        return False, "Please enter a more specific question."
    
    # Check if query contains only special characters or numbers
    import re
    if re.match(r'^[^a-zA-Z]*$', query_stripped):
        return False, "Please enter a question using words."
    
    # Check for only repeated characters
    if len(set(query_stripped.lower())) <= 2:
        return False, "Please enter a meaningful question."
    
    # Check for test inputs
    test_patterns = [
        r'^test+$', r'^hello+$', r'^hi+$', r'^\.+$', r'^\?+$', 
        r'^!+$', r'^@+$', r'^#+$', r'^\$+$', r'^%+$', r'^\^+$',
        r'^&+$', r'^\*+$', r'^\(+$', r'^\)+$', r'^-+$', r'^_+$',
        r'^=+$', r'^\++$', r'^\[+$', r'^\]+$', r'^\{+$', r'^\}+$',
        r'^\\+$', r'^\|+$', r'^;+$', r'^:+$', r'^"+$', r"^'+$",
        r'^<+$', r'^>+$', r'^,+$', r'^\.+$', r'^/+$', r'^\s*$'
    ]
    
    for pattern in test_patterns:
        if re.match(pattern, query_stripped.lower()):
            return False, "Please enter a meaningful compliance-related question."
    
    # Check for minimum word count
    words = query_stripped.split()
    if len(words) < 2:
        return False, "Please enter a more detailed question."
    
    return True, ""

def detect_conversation_history_query(query: str) -> bool:
    """Detect if user is asking about previous conversation."""
    # More specific patterns to avoid false positives
    specific_history_patterns = [
        "what did we talk about before",
        "what did we talk about earlier", 
        "what was our last conversation",
        "what was our previous conversation",
        "what did we discuss before",
        "what did we discuss earlier",
        "our chat history",
        "previous chat",
        "last conversation",
        "continue our conversation",
        "where did we leave off",
        "what were we discussing",
        "remind me what we talked about"
    ]
    
    query_lower = query.lower().strip()
    
    # Exact matches for specific patterns
    for pattern in specific_history_patterns:
        if pattern in query_lower:
            return True
    
    # Additional check for very short history queries
    if len(query_lower.split()) <= 6:  # Short queries only
        history_keywords = ["before", "earlier", "last time", "previous"]
        if any(keyword in query_lower for keyword in history_keywords):
            # Make sure it's actually about conversation
            conversation_words = ["talk", "discuss", "chat", "conversation", "said", "spoke"]
            if any(word in query_lower for word in conversation_words):
                return True
    
    return False

def handle_conversation_history_query(conversation: 'ConversationHistory') -> str:
    """Handle queries about conversation history."""
    if not conversation or not conversation.history:
        return """I don't have any record of previous conversations in this session. This appears to be a new conversation.

If you're looking for help with compliance topics, I'd be happy to assist with:
• GDPR, CCPA, HIPAA compliance
• Security frameworks (ISO 27001, SOC 2, NIST)
• Privacy policy development
• Risk assessments and audits

What specific compliance topic would you like to discuss?"""
    
    # If there is history, provide a summary
    context = conversation.get_context(compact=True)
    return f"""Here's a summary of our previous conversation:

{context}

What would you like to continue discussing or explore further?"""

def rate_limited_generate_content(prompt: str, temperature: float = 0.1, max_tokens: int = 3200, max_retries: int = 2) -> str:
    """Generate content with rate limiting and retries."""
    if not _ensure_model_initialized():
        logger.error("Gemini model unavailable; returning empty response")
        return ""

    # Check cache first for small prompts (classification, guardrails, etc.)
    if max_tokens <= 300:
        prompt_hash = hash_text(f"{prompt}:{temperature}:{max_tokens}")
        cache_key = f"gemini_small:{prompt_hash}"
        if cache_key in QUERY_CACHE:
            cached = QUERY_CACHE[cache_key]
            if cached and len(cached.strip()) > 0:
                logger.info(f"✅ Cache hit for small prompt (tokens: {max_tokens})")
                return cached

    last_key_index = _ACTIVE_GEMINI_INDEX
    for attempt in range(max_retries):
        try:
            _log_gemini_api_call(
                context="rate_limited_generate_content",
                prompt=prompt,
                temperature=temperature,
                max_tokens=max_tokens,
                optimized=False,
            )
            response = client.models.generate_content(
                model=MODEL_ID,
                contents=prompt,
                config=types.GenerateContentConfig(
                    thinking_config=types.ThinkingConfig(thinking_level="low"),
                    temperature=temperature,
                    max_output_tokens=max_tokens,
                )
            )
            
            if response and response.text:
                result_text = response.text.strip()
                # Cache small responses
                if max_tokens <= 300 and result_text:
                    prompt_hash = hash_text(f"{prompt}:{temperature}:{max_tokens}")
                    cache_key = f"gemini_small:{prompt_hash}"
                    QUERY_CACHE[cache_key] = result_text
                return result_text
            else:
                logger.warning("Empty response from Gemini API")
                if attempt < max_retries - 1:
                    time.sleep(2 ** attempt)  # Exponential backoff: 1s, 2s
                continue
                
        except Exception as e:
            error_str = str(e).lower()
            error_message = str(e)
            is_rate_limit = any(k in error_str for k in ["429", "resourceexhausted", "quota", "rate limit", "rate_limit"])
            is_quota_exhausted = _is_quota_exhausted(error_message)
            is_key_invalid = _is_key_leaked_or_invalid(error_message)
            
            # Mark key as invalid if leaked/revoked
            if is_key_invalid and _ACTIVE_GEMINI_INDEX is not None:
                _mark_key_invalid(_ACTIVE_GEMINI_INDEX, error_message[:200])
                # Don't retry invalid keys
                if attempt == max_retries - 1:
                    logger.error(f"Key #{_ACTIVE_GEMINI_INDEX + 1} is invalid. All retries exhausted.")
                    break
                # Try rotating to another key immediately
                if _try_all_keys_rotation() and _ACTIVE_GEMINI_INDEX != last_key_index:
                    logger.info(f"Rotated away from invalid key #{last_key_index + 1} to key #{_ACTIVE_GEMINI_INDEX + 1}")
                    last_key_index = _ACTIVE_GEMINI_INDEX
                    continue
                else:
                    logger.error("No valid keys available. Invalid key cannot be used.")
                    break
            
            if attempt == max_retries - 1:
                logger.error(f"All retries exhausted. Last error: {e}")
                break
            
            # Extract retry delay from API response
            retry_delay = _extract_retry_delay(error_message)
            
            # Exponential backoff for rate limits
            if is_rate_limit or is_quota_exhausted:
                # Mark current key as exhausted if quota exhausted
                if is_quota_exhausted and _ACTIVE_GEMINI_INDEX is not None:
                    _mark_key_exhausted(_ACTIVE_GEMINI_INDEX)
                    logger.error(f"Key #{_ACTIVE_GEMINI_INDEX + 1} quota exhausted. Marked for cooldown.")
                
                # Use API-provided retry delay if available, otherwise use exponential backoff
                if retry_delay:
                    backoff_time = min(retry_delay, 120)  # Cap at 2 minutes
                    logger.warning(f"Quota/Rate limit detected on key #{_ACTIVE_GEMINI_INDEX + 1 if _ACTIVE_GEMINI_INDEX is not None else 'unknown'}. API suggests retry in {retry_delay}s. Waiting {backoff_time}s...")
                else:
                    backoff_time = min(5 * (2 ** attempt), 60)  # Max 60s for exponential backoff
                    logger.warning(f"Rate limit detected on key #{_ACTIVE_GEMINI_INDEX + 1 if _ACTIVE_GEMINI_INDEX is not None else 'unknown'} (attempt {attempt+1}/{max_retries}), waiting {backoff_time}s before retry")
                
                # Try switching to a different key (even for quota exhaustion, as keys might be from different accounts)
                if _try_all_keys_rotation() and _ACTIVE_GEMINI_INDEX != last_key_index:
                    logger.info(f"Rotated to key #{_ACTIVE_GEMINI_INDEX + 1} after quota/rate limit. Retrying immediately...")
                    last_key_index = _ACTIVE_GEMINI_INDEX
                    # Don't wait if we successfully rotated to a new key
                    continue
                else:
                    # All keys exhausted or no other keys available
                    if is_quota_exhausted:
                        logger.error(f"All keys appear exhausted. Waiting {backoff_time}s before retry...")
                    else:
                        logger.warning(f"All keys rate limited. Waiting {backoff_time}s before retry...")
                    time.sleep(backoff_time)
                    continue
            else:
                # Non-rate-limit error - try switching key
                if _switch_to_fallback_key() and _ACTIVE_GEMINI_INDEX != last_key_index:
                    logger.info(f"Switched to fallback Gemini API key #{_ACTIVE_GEMINI_INDEX + 1} after error: {str(e)[:100]}")
                    last_key_index = _ACTIVE_GEMINI_INDEX
                    time.sleep(1)  # Short delay before retry
                    continue
                else:
                    # Wait before retrying
                    time.sleep(2 ** attempt)  # Exponential backoff: 1s, 2s
    
    # All retries failed - return empty response
    logger.error("All Gemini API attempts failed. Quota may be exhausted.")
    return ""

def detect_document_analysis_request(query: str) -> bool:
    """Detect if user is asking to analyze their uploaded document.

    Requires both an analysis verb and a document reference to avoid
    false positives on general compliance questions.
    """
    try:
        if not isinstance(query, str):
            return False
        q = query.lower()
        analysis_verbs = [
            "analyze", "analyse", "review", "check", "assess", "evaluate", "summarize", "read", "tell me about"
        ]
        document_terms = [
            "document", "file", "pdf", "docx", "policy", "privacy policy", "terms", "terms and conditions",
            "this", "it"  # Added to catch "analyze this" or "check it"
        ]
        has_analysis = any(v in q for v in analysis_verbs)
        has_doc_ref = any(t in q for t in document_terms)
        return has_analysis and has_doc_ref
    except Exception:
        return False

def detect_document_reference(query: str, conversation_context: str = "") -> bool:
    """Detect whether the user is referring to a specific uploaded document using intelligent pattern matching.

    Uses multiple heuristics to identify document references without relying on simple keyword matching:
    1. Strong possession patterns (my/the/this doc I uploaded, file I provided)
    2. Upload/possession indicators combined with compliance questions
    3. Explicit self-reference with compliance context
    4. Context-aware patterns (looks at conversation history for document mentions)
    
    IMPORTANT: Primarily focuses on the CURRENT QUERY to avoid false positives from conversation history.
    """
    try:
        query_lower = (query or "").lower().strip()
        if not query_lower:
            return False

        # Pattern 1: Strong possession indicators with upload context (CURRENT QUERY ONLY)
        strong_possession_patterns = [
            r"\b(?:my|this|our)\s+(?:doc|document|file|pdf|docx|policy|privacy\s+policy|terms)\s+(?:i|we|that\s+i)\s+(?:uploaded|provided|attached|sent|have)",
            r"\b(?:uploaded|provided|attached|sent)\s+(?:my|this|our)\s+(?:doc|document|file|pdf|docx|policy|privacy\s+policy|terms)",
            r"\b(?:the\s+)?(?:doc|document|file)\s+(?:i|we)\s+(?:uploaded|provided|attached|sent|have)\s+(?:is|for|to)",
            r"\b(?:check|review|analyze|examine|assess)\s+(?:my|this|our)\s+(?:uploaded|provided|attached|sent)?\s+(?:doc|document|file|policy)",
            r"\b(?:is|does)\s+(?:my|this)\s+(?:uploaded|provided)\s+(?:doc|document|file|policy)\s+(?:compliant|comply|follow|meet)",
            r"\b(?:analyze|check|review)\s+(?:the\s+)?(?:doc|document|file|policy)(?:\s+(?:i|we)\s+uploaded)?",
            r"\b(?:uploaded|attached|provided)\s+(?:doc|document|file|policy)"
        ]

        # Pattern 2: Self-reference with specific compliance questions (CURRENT QUERY ONLY)
        self_reference_patterns = [
            r"\b(?:my|this)\s+(?:doc|document|file|policy|privacy\s+policy|terms)\s+(?:is|does|can|will|should)\s+(?:compliant|comply|follow|meet|according\s+to|align\s+with)",
            r"\b(?:is|does|can)\s+(?:my|this)\s+(?:doc|document|file|policy)\s+(?:gdpr|hipaa|ccpa|compliant|comply)",
            r"\b(?:tell\s+me\s+about|what\s+about|how\s+about)\s+(?:my|this)\s+(?:uploaded|provided)\s+(?:doc|document|file|policy)"
        ]

        # Pattern 3: Explicit upload/file references (CURRENT QUERY ONLY)
        explicit_document_patterns = [
            r"\b(?:doc|document|file|policy)\s+(?:i|we)\s+(?:uploaded|provided|have|sent)\s+(?:for|to|with)",
            r"\b(?:uploaded|provided)\s+(?:doc|document|file|policy)",
            r"\b(?:the\s+)?(?:file|document)\s+(?:i\s+)?(?:uploaded|attached|sent|provided)"
        ]

        # Pattern 4: General/informational question patterns (should NOT trigger)
        general_question_patterns = [
            r"^(?:what|who|when|where|why|how)\s+(?:is|are|does|do|can|should|would)",
            r"\b(?:tell\s+me\s+about|explain|describe|define)\s+(?!my|this|the\s+uploaded)",
            r"\b(?:what\s+should|how\s+do\s+i|can\s+you\s+explain)",
            r"\b(?:if\s+i\s+had|suppose\s+i\s+have|what\s+would|how\s+would\s+you)",
            r"\b(?:in\s+general|generally|typically|usually|standard|best\s+practice)",
            r"\b(?:example|sample|template)"
        ]

        # First check: Is this clearly a general question? If yes, return False immediately
        for pattern in general_question_patterns:
            if re.search(pattern, query_lower):
                return False

        # Second check: Look for explicit document references in CURRENT QUERY ONLY
        for pattern in strong_possession_patterns + self_reference_patterns + explicit_document_patterns:
            if re.search(pattern, query_lower):
                return True

        # Third check: Only use conversation context if the current query has some document-related words
        # This prevents "what is gdpr" from being flagged just because previous message mentioned "document"
        doc_keywords_in_query = re.search(r"\b(?:doc|document|file|policy|uploaded|attached|provided|sent)\b", query_lower)
        
        if doc_keywords_in_query and conversation_context:
            context_lower = conversation_context.lower()
            # Only flag as document reference if BOTH query has doc keywords AND context shows upload activity
            upload_in_context = re.search(r"\b(?:uploaded|provided|attached|sent|upload)\b", context_lower)
            if upload_in_context:
                return True

        return False
    except Exception:
        return False

def analyze_general_documentation_compliance(document_text: str, frameworks: List[str]) -> str:
    """Analyze general system/technical documentation for compliance issues against specified frameworks.
    
    Args:
        document_text: The content of the system documentation
        frameworks: List of frameworks to check against (e.g., ['GDPR', 'ISO 27001', 'SOC 2'])
    
    Returns:
        Comprehensive compliance analysis report
    """
    try:
        frameworks_str = ", ".join(frameworks) if frameworks else "GDPR, ISO 27001, SOC 2, HIPAA"
        
        prompt = f"""You are a compliance security expert. Analyze the following system/technical documentation for compliance issues and security gaps against these frameworks: {frameworks_str}

DOCUMENTATION:
{document_text[:8000]}

Provide a comprehensive compliance analysis with the following structure:

## 📋 Document Overview
- Brief summary of what the document describes
- Key systems/components identified

## 🔍 Compliance Analysis per Framework

For each framework ({frameworks_str}):

### [Framework Name] Compliance Assessment

**✅ Compliant Areas:**
- List aspects that meet framework requirements
- Reference specific sections/controls

**❌ Gaps & Issues:**
- Identify missing controls or requirements
- Highlight security concerns
- Rate severity (Critical/High/Medium/Low)

**📊 Compliance Score:** [X/10]

## 🔐 Security & Privacy Findings

**Data Protection:**
- How is data handled, stored, encrypted?
- Are there data retention/deletion policies?
- Is access control properly defined?

**Authentication & Authorization:**
- How are users authenticated?
- What authorization mechanisms exist?
- Are there privilege escalation risks?

**Audit & Monitoring:**
- Are there audit logs?
- What monitoring is in place?
- Incident response procedures?

**Infrastructure Security:**
- Network security measures?
- Encryption in transit/at rest?
- Backup and disaster recovery?

## ⚠️ Critical Risks Identified
1. [Risk description] - [Impact] - [Recommendation]
2. [Risk description] - [Impact] - [Recommendation]

## ✨ Recommendations

**Immediate Actions (Priority: High):**
1. [Specific action]
2. [Specific action]

**Short-term Improvements (1-3 months):**
1. [Specific action]
2. [Specific action]

**Long-term Enhancements (3-6 months):**
1. [Specific action]
2. [Specific action]

## 📝 Missing Documentation
List any critical documentation that should exist but is missing:
- [ ] Data flow diagrams
- [ ] Security architecture
- [ ] Incident response plan
- [ ] etc.

Provide specific, actionable insights with clear references to framework requirements."""

        analysis = rate_limited_generate_content_optimized(prompt, temperature=0.2, max_tokens=4000)
        return analysis
        
    except Exception as e:
        logger.error(f"Error analyzing general documentation: {e}")
        return f"Error analyzing documentation: {str(e)}"

def analyze_query_intent_with_ai(query: str, conversation_context: str = "") -> Dict[str, Any]:
    """Analyze query intent using AI, considering conversation context when provided."""
    try:
        prompt = f"""Analyze the following query and determine its intent.

Query: {query}
Conversation Context: {conversation_context}

You must respond with ONLY a valid JSON object, nothing else. No explanations, no markdown, just pure JSON.

Return this exact JSON structure:
{{
  "intent": "GENERAL_COMPLIANCE",
  "document_type": "unknown",
  "framework": "general",
  "urgency": "medium",
  "confidence": 0.5,
  "reasoning": "Brief explanation"
}}

Valid intent values: GENERAL_COMPLIANCE, SPECIFIC_REQUIREMENT, DOCUMENT_ANALYSIS, ANALYZE_UPLOADED
Valid urgency values: low, medium, high
Confidence must be between 0.0 and 1.0"""
        
        response = rate_limited_generate_content(prompt, temperature=0.1, max_tokens=200)
        
        # Clean response - remove markdown code blocks if present
        cleaned = response.strip()
        if cleaned.startswith('```'):
            # Remove markdown code block formatting
            lines = cleaned.split('\n')
            cleaned = '\n'.join([l for l in lines if not l.strip().startswith('```')])
            cleaned = cleaned.strip()
        
        # Try to extract JSON if embedded in text
        if '{' in cleaned and '}' in cleaned:
            start = cleaned.find('{')
            end = cleaned.rfind('}') + 1
            cleaned = cleaned[start:end]
        
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        logger.warning(f"AI classification failed - JSON parse error: {str(e)}")
        return {
            'intent': 'GENERAL_COMPLIANCE',
            'document_type': 'unknown',
            'framework': 'general',
            'urgency': 'medium',
            'confidence': 0.0,
            'reasoning': f'Default classification due to JSON error'
        }
    except Exception as e:
        logger.warning(f"AI classification failed: {str(e)}")
        return {
            'intent': 'GENERAL_COMPLIANCE',
            'document_type': 'unknown',
            'framework': 'general',
            'urgency': 'medium',
            'confidence': 0.0,
            'reasoning': f'Default classification due to error'
        }

if __name__ == "__main__":
    main() 